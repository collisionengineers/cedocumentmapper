import json
import os
import re
import sys
import io
import shutil
import subprocess
import tempfile
import zipfile
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from xml.etree import ElementTree as ET
from email import policy
from email.parser import BytesParser

import tkinter as tk
from tkinter import messagebox, ttk

try:
    from tkinterdnd2 import DND_FILES, TkinterDnD  # optional
    HAS_DND = True
except Exception:
    HAS_DND = False
    DND_FILES = None
    TkinterDnD = None

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from pypdf import PdfReader

try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None

try:
    import extract_msg  # pure-Python .msg parser
except Exception:
    extract_msg = None

try:
    import pytesseract  # OCR fallback for scanned-only PDFs
except Exception:
    pytesseract = None

try:
    from PIL import Image as _PILImage  # used by the OCR fallback
except Exception:
    _PILImage = None


APP_TITLE = "CE Document Mapper"
MINIMAL_WINDOW_GEOMETRY = "400x720"
EXPANDED_WINDOW_GEOMETRY = "1400x800"
RJS_PROVIDER_NAME = "RJS"
RJS_EXPORT_BASENAME = "RJS"

# Border color used on a value entry in the left panel when an Engineer
# Report has overwritten that field's value. The highlight is permanent
# until a new document is dragged in.
ENGINEER_OVERRIDE_BORDER_COLOR = "#d62828"

DEFAULT_FIELDS = [
    ("work_provider", "Work Provider"),
    ("vrm", "VRM"),
    ("vehicle_model", "Vehicle Model"),
    ("claimant_name", "Claimant Name"),
    ("reference", "Reference"),
    ("incident_date", "Incident Date"),
    ("instruction_date", "Instruction Date"),
    ("inspection_date", "Inspection Date"),
    ("inspection_address", "Inspection Address"),
    ("accident_circumstances", "Accident Circumstances"),
    ("vat_status", "VAT Status"),
    ("mileage", "Mileage"),
    ("mileage_unit", "Mileage Unit"),
]

REQUIRED_FIELDS = {
    "work_provider",
    "vrm",
    "vehicle_model",
    "claimant_name",
    "reference",
    "incident_date",
    "instruction_date",
}

FIELD_LABELS = dict(DEFAULT_FIELDS)
FIELD_KEYS = [key for key, _ in DEFAULT_FIELDS]
NON_PROVIDER_FIELDS = [key for key in FIELD_KEYS if key != "work_provider"]

METHOD_CHOICES = [
    ("single_label", "Single Label"),
    ("two_labels", "Two Labels"),
    ("fixed_position", "Fixed Position"),
    ("fixed_position_label", "Fixed Position + Label"),
    ("single_label_offset", "Single Label +/-"),
    ("email_date", "Email Date"),
    ("manual_input", "Manual Input"),
]
METHOD_LABEL_TO_CODE = {label: code for code, label in METHOD_CHOICES}
METHOD_CODE_TO_LABEL = {code: label for code, label in METHOD_CHOICES}
LEGACY_METHOD_TO_DISPLAY_CODE = {
    "labels": "single_label",
    "multiline_labels": "two_labels",
    "letterhead_date": "fixed_position",
    "date_near_keywords": "two_labels",
    "reference_fallback": "fixed_position",
    "vrm_fallback": "fixed_position",
    "claimant_name_fallback": "fixed_position",
    "vehicle_model_fallback": "fixed_position",
    "address_fallback": "fixed_position",
    "fixed_value": "manual_input",
    "current_date": "manual_input",
    "blank": "manual_input",
}
DEFAULT_METHOD_BY_FIELD = {
    "vrm": "single_label",
    "vehicle_model": "single_label",
    "claimant_name": "single_label",
    "reference": "single_label",
    "incident_date": "single_label",
    "instruction_date": "two_labels",
    "inspection_date": "manual_input",
    "inspection_address": "two_labels",
    "mileage": "single_label",
    "mileage_unit": "single_label",
    "accident_circumstances": "two_labels",
    "vat_status": "single_label",
}

# Fields whose mapping is a simple "look for these tokens anywhere in the
# document" presence check rather than a labelled extraction. The user enters
# tokens describing the positive scenario; if any token is found in the
# document the field is set to ``positive_value``, otherwise to
# ``negative_value``. A blank config leaves the field blank.
#
# These fields hide the Method dropdown in the provider editor since the
# method choice is irrelevant — only the configured tokens matter.
PRESENCE_CHECK_FIELDS = {
    "vat_status": {"positive_value": "Yes", "negative_value": "No"},
    "mileage_unit": {"positive_value": "Miles", "negative_value": "Km"},
}

# Cap on how many pages we'll OCR. Real scanned-letter instructions
# observed in practice are 1-2 pages. Image-dump PDFs (photo evidence
# bundles) start at 3+ pages and would silently waste 60+ seconds of
# OCR time before the app responded again. Any document above this
# limit skips OCR entirely.
OCR_PAGE_LIMIT = 2


def resource_path(name: str) -> Path:
    if hasattr(sys, "_MEIPASS"):
        return Path(getattr(sys, "_MEIPASS")) / name
    return Path(__file__).resolve().parent / name


def configure_bundled_tesseract() -> bool:
    """Point ``pytesseract`` at the Tesseract bundled with the .exe.

    Looks for ``tesseract/tesseract.exe`` (Windows) or
    ``tesseract/tesseract`` (other) under PyInstaller's ``sys._MEIPASS``
    or next to ``app.py`` in development. Also sets the ``TESSDATA_PREFIX``
    environment variable so Tesseract can find its language data even
    when running from the bundle.

    Returns True if a working Tesseract binary was located, False
    otherwise. ``pytesseract`` itself does not need to be initialised
    until first use; this function only sets up the paths.
    """
    if pytesseract is None:
        return False
    base = resource_path("tesseract")
    if not base.exists():
        return False
    # Locate the binary
    candidates = [
        base / "tesseract.exe",
        base / "tesseract",
    ]
    binary = next((c for c in candidates if c.exists()), None)
    if binary is None:
        return False
    try:
        pytesseract.pytesseract.tesseract_cmd = str(binary)
    except Exception:
        return False
    # The trained-data folder. Tesseract looks for ``tessdata`` next to
    # the binary by default; set the env var explicitly to be safe.
    tessdata = base / "tessdata"
    if tessdata.exists():
        os.environ.setdefault("TESSDATA_PREFIX", str(tessdata))
    return True


def apply_window_icon(root: tk.Misc) -> None:
    icon_path = resource_path("ce_document_mapper.ico")
    if not icon_path.exists():
        return
    try:
        root.iconbitmap(str(icon_path))
    except Exception:
        try:
            root.wm_iconbitmap(str(icon_path))
        except Exception:
            pass


def _windows_known_folder_path(folder_guid: str) -> Optional[Path]:
    """Resolve a Windows ``KNOWNFOLDERID`` to its actual current path.

    Uses ``SHGetKnownFolderPath`` from the Windows Shell API. This is
    the only reliable way to find the user's real Desktop / Documents
    folder on machines with OneDrive (or any other) folder
    redirection — ``Path.home() / "Desktop"`` returns the legacy
    folder, which on a redirected machine is empty and invisible to
    the user.

    Returns ``None`` on non-Windows platforms or if the Shell call
    fails for any reason; callers fall back to the home-relative path.
    """
    if sys.platform != "win32":
        return None
    try:
        import ctypes
        from ctypes import wintypes

        # SHGetKnownFolderPath signature:
        #   HRESULT SHGetKnownFolderPath(
        #     REFKNOWNFOLDERID rfid, DWORD dwFlags,
        #     HANDLE hToken, PWSTR *ppszPath)
        # We pass a GUID, get back a wide-string path.
        class _GUID(ctypes.Structure):
            _fields_ = [
                ("Data1", wintypes.DWORD),
                ("Data2", wintypes.WORD),
                ("Data3", wintypes.WORD),
                ("Data4", ctypes.c_ubyte * 8),
            ]

        # Parse "{XXXXXXXX-XXXX-XXXX-XXXX-XXXXXXXXXXXX}" into _GUID.
        clean = folder_guid.strip("{}")
        parts = clean.split("-")
        if len(parts) != 5:
            return None
        guid = _GUID()
        guid.Data1 = int(parts[0], 16)
        guid.Data2 = int(parts[1], 16)
        guid.Data3 = int(parts[2], 16)
        rest = bytes.fromhex(parts[3] + parts[4])
        for i, byte in enumerate(rest):
            guid.Data4[i] = byte

        SHGetKnownFolderPath = ctypes.windll.shell32.SHGetKnownFolderPath
        SHGetKnownFolderPath.argtypes = [
            ctypes.POINTER(_GUID),
            wintypes.DWORD,
            wintypes.HANDLE,
            ctypes.POINTER(ctypes.c_wchar_p),
        ]
        SHGetKnownFolderPath.restype = ctypes.HRESULT

        out_ptr = ctypes.c_wchar_p()
        result = SHGetKnownFolderPath(ctypes.byref(guid), 0, None, ctypes.byref(out_ptr))
        if result != 0 or not out_ptr.value:
            return None
        path = Path(out_ptr.value)
        # Free the buffer Windows allocated for us.
        ctypes.windll.ole32.CoTaskMemFree(out_ptr)
        return path if path.exists() else None
    except Exception:
        return None


# Standard KNOWNFOLDERID GUIDs from <KnownFolders.h>:
_FOLDERID_DESKTOP = "{B4BFCC3A-DB2C-424C-B029-7FE99A87C641}"
_FOLDERID_DOCUMENTS = "{FDD39AD0-238F-46AF-ADB4-6C85480369C7}"


def get_documents_dir() -> Path:
    # On Windows, ask the Shell where Documents really is so we honour
    # OneDrive redirection. Fall back to the home-relative path on
    # other platforms or if the Shell call fails.
    resolved = _windows_known_folder_path(_FOLDERID_DOCUMENTS)
    if resolved is not None:
        return resolved
    home = Path.home()
    docs = home / "Documents"
    return docs if docs.exists() else home


def get_desktop_dir() -> Path:
    # Same story as Documents — ask the Shell for the real Desktop on
    # Windows so OneDrive-redirected users see their exports where they
    # expect.
    resolved = _windows_known_folder_path(_FOLDERID_DESKTOP)
    if resolved is not None:
        return resolved
    home = Path.home()
    desktop = home / "Desktop"
    return desktop if desktop.exists() else home


def _migrate_legacy_app_data(new_dir: Path) -> None:
    """One-time migration of ``providers.json`` and ``app_settings.json``.

    Earlier versions resolved the user's Documents folder via
    ``Path.home() / "Documents"``, which on OneDrive-redirected
    machines points at the legacy (empty, invisible) folder. V62
    switched to the Shell-resolved path. For users upgrading, copy
    any existing files from the legacy location into the new one so
    custom presets and saved settings aren't silently abandoned.

    Runs once: if the new location already has a file, the legacy
    one is left untouched (the new location wins).
    """
    if sys.platform != "win32":
        return
    legacy_home_documents = Path.home() / "Documents"
    if not legacy_home_documents.exists():
        return
    legacy_app_dir = legacy_home_documents / APP_TITLE
    if not legacy_app_dir.exists() or legacy_app_dir.resolve() == new_dir.resolve():
        return
    new_dir.mkdir(parents=True, exist_ok=True)
    for filename in ("providers.json", "app_settings.json"):
        legacy_file = legacy_app_dir / filename
        new_file = new_dir / filename
        if legacy_file.exists() and not new_file.exists():
            try:
                shutil.copy2(legacy_file, new_file)
            except Exception:
                # Best-effort migration. If copying fails the user
                # gets a fresh seed from the bundled providers.json,
                # which is acceptable.
                pass


APP_DATA_DIR = get_documents_dir() / APP_TITLE
# Migrate any legacy provider/settings files from the pre-V62 location
# (the Path.home()-relative Documents folder) into the new
# Shell-resolved location, so OneDrive-redirected users keep their
# customisations across the upgrade.
_migrate_legacy_app_data(APP_DATA_DIR)
APP_DATA_DIR.mkdir(parents=True, exist_ok=True)

SETTINGS_PATH = APP_DATA_DIR / "app_settings.json"
DEFAULT_CONFIG_PATH = APP_DATA_DIR / "providers.json"
OUTPUT_DIR = get_desktop_dir()

DEFAULT_CONFIG = {
    "providers": [
        {
            "name": RJS_PROVIDER_NAME,
            "detect_phrases": [
                "robert james solicitors",
                "urgent vehicle inspection required",
                "our client:",
            ],
            "field_rules": {
                "vrm": {"method": "single_label", "config": "Client vehicle registration, Vehicle Registration, Registration"},
                "vehicle_model": {"method": "single_label", "config": "Client vehicle model, Vehicle Model"},
                "claimant_name": {"method": "single_label", "config": "Our Client, Client"},
                "reference": {"method": "single_label", "config": "Our Reference, Reference"},
                "incident_date": {"method": "single_label", "config": "Accident, Date of Accident, Accident Date"},
                "instruction_date": {"method": "two_labels", "config": "Email: || URGENT VEHICLE INSPECTION REQUIRED"},
                "inspection_date": {"method": "manual_input", "config": ""},
                "inspection_address": {"method": "two_labels", "config": "Address: || Mobile Tel:"},
                "mileage": {"method": "manual_input", "config": ""},
                "mileage_unit": {"method": "manual_input", "config": ""},
                "accident_circumstances": {"method": "two_labels", "config": "The circumstances of the accident are || Please arrange an inspection"},
                "vat_status": {"method": "manual_input", "config": ""},
            },
            "use_current_date_for_inspection_date": False,
            "force_postcode_for_inspection_address": False,
            "engineer_report": False,
        }
    ]
}


@dataclass
class ProviderMatch:
    name: str
    score: int
    config: dict


@dataclass
class DocumentSession:
    path: str
    text: str
    provider_match: ProviderMatch
    values: Dict[str, str]
    notes: List[str]
    selected_provider: str = ""
    detected_provider_name: str = ""
    source_paths: Optional[List[str]] = None


class ExtractionEngine:
    def __init__(self, config_path: str):
        self.config_path = str(config_path)
        self.config = self.load_or_create_config(self.config_path)

    def load_or_create_config(self, config_path: str) -> dict:
        path = Path(config_path)
        if not path.exists():
            path.parent.mkdir(parents=True, exist_ok=True)
            # If a ``providers.json`` was bundled alongside the .exe (or
            # next to ``app.py`` in development), use it to seed the
            # user's first run. Otherwise fall back to the hardcoded
            # ``DEFAULT_CONFIG``. The seeding only runs when the user
            # has no existing providers.json — anyone who already has
            # a file in Documents is left alone.
            seed_payload: Optional[dict] = None
            try:
                bundled_path = resource_path("providers.json")
                if bundled_path.exists():
                    bundled_data = json.loads(bundled_path.read_text(encoding="utf-8"))
                    bundled_data.setdefault("providers", [])
                    seed_payload = bundled_data
            except Exception:
                # Anything unexpected with the bundled file -> silent
                # fall-through to the hardcoded default.
                seed_payload = None
            if seed_payload is None:
                seed_payload = json.loads(json.dumps(DEFAULT_CONFIG))
            path.write_text(
                json.dumps(seed_payload, indent=2, ensure_ascii=False),
                encoding="utf-8",
            )
        data = json.loads(path.read_text(encoding="utf-8"))
        data.setdefault("providers", [])
        data["providers"] = [self.normalize_provider_config(provider) for provider in data.get("providers", [])]
        return data

    def normalize_provider_config(self, provider: dict) -> dict:
        provider = dict(provider or {})
        provider.setdefault("name", "")
        provider.setdefault("detect_phrases", [])
        provider.setdefault("use_current_date_for_inspection_date", False)
        provider.setdefault("force_postcode_for_inspection_address", False)
        provider.setdefault("engineer_report", False)

        field_rules = provider.get("field_rules") or {}
        legacy_labels = provider.get("field_labels") or {}

        # Work Provider is a manual-input-only field on the preset. If
        # the preset doesn't have one yet (e.g. older presets created
        # before V56), default to a blank rule so the user is forced to
        # fill it in next time they edit the preset.
        wp_rule = dict(field_rules.get("work_provider") or {})
        wp_rule.setdefault("method", "manual_input")
        wp_rule.setdefault("config", "")
        field_rules["work_provider"] = wp_rule

        for field_name in NON_PROVIDER_FIELDS:
            rule = dict(field_rules.get(field_name) or {})
            method = rule.get("method")
            config = rule.get("config", "")

            if not method:
                if field_name == "instruction_date":
                    method = "two_labels"
                    config = config or ""
                elif field_name == "inspection_address":
                    method = "two_labels"
                    labels = legacy_labels.get(field_name, [])
                    start_label = labels[0] if labels else "Address"
                    config = f"{start_label} || Mobile Tel"
                else:
                    method = DEFAULT_METHOD_BY_FIELD.get(field_name, "single_label")
                    config = ", ".join(legacy_labels.get(field_name, []))

            if method in LEGACY_METHOD_TO_DISPLAY_CODE:
                display_code = LEGACY_METHOD_TO_DISPLAY_CODE[method]
                if method == "fixed_value":
                    method = "manual_input"
                elif method == "current_date":
                    method = "manual_input"
                    config = "{today}"
                elif method == "blank":
                    method = "manual_input"
                    config = ""
                elif method == "multiline_labels" and "||" not in config:
                    labels = [part.strip() for part in config.split(",") if part.strip()]
                    start_label = labels[0] if labels else "Address"
                    method = "two_labels"
                    config = f"{start_label} || Mobile Tel"
                elif method == "labels":
                    method = "single_label"
                else:
                    method = display_code

            field_rules[field_name] = {"method": method, "config": config or ""}

        provider["field_rules"] = field_rules
        provider.pop("field_labels", None)
        provider.pop("use_image_based_assessment_for_inspection_address", None)
        provider.pop("extract_images_on_export", None)
        return provider

    def save_config(self, config: Optional[dict] = None):
        if config is not None:
            self.config = config
        target = Path(self.config_path)
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_text(json.dumps(self.config, indent=2, ensure_ascii=False), encoding="utf-8")

    def list_provider_names(self) -> List[str]:
        return sorted(
            [p.get("name", "").strip() for p in self.config.get("providers", []) if p.get("name", "").strip()],
            key=lambda x: x.lower(),
        )

    def get_provider_config(self, name: str) -> Optional[dict]:
        for provider in self.config.get("providers", []):
            if provider.get("name", "").strip().lower() == name.strip().lower():
                return provider
        return None

    def upsert_provider(self, provider_data: dict):
        providers = self.config.setdefault("providers", [])
        provider_data = self.normalize_provider_config(provider_data)
        name = provider_data.get("name", "").strip()
        if not name:
            raise ValueError("Provider name is required.")
        for idx, provider in enumerate(providers):
            if provider.get("name", "").strip().lower() == name.lower():
                providers[idx] = provider_data
                self.save_config()
                return
        providers.append(provider_data)
        providers.sort(key=lambda x: x.get("name", "").lower())
        self.save_config()

    @staticmethod
    def _read_text_file(path: Path) -> str:
        for encoding in ("utf-8", "cp1252", "latin-1"):
            try:
                return path.read_text(encoding=encoding)
            except Exception:
                continue
        return path.read_text(errors="ignore")


    def _extract_docx_textbox_lines(self, path: str) -> List[str]:
        textbox_lines: List[str] = []
        try:
            with zipfile.ZipFile(path) as zf:
                xml_names = [
                    name for name in zf.namelist()
                    if name.startswith("word/") and name.endswith(".xml")
                ]
                for name in xml_names:
                    try:
                        root = ET.fromstring(zf.read(name))
                    except Exception:
                        continue

                    for node in root.iter():
                        tag_name = node.tag.rsplit("}", 1)[-1].lower()
                        if tag_name not in {"txbxcontent", "textbox"}:
                            continue

                        collected: List[str] = []
                        for child in node.iter():
                            child_tag = child.tag.rsplit("}", 1)[-1].lower()
                            if child_tag == "t":
                                txt = (child.text or "").strip()
                                if txt:
                                    collected.append(txt)
                            elif child_tag in {"p", "br", "cr"} and collected:
                                textbox_lines.append(" ".join(collected).strip())
                                collected = []
                        if collected:
                            textbox_lines.append(" ".join(collected).strip())
        except Exception:
            return []

        cleaned: List[str] = []
        seen = set()
        for line in textbox_lines:
            line = self.clean_value(line)
            if line and line.lower() not in seen:
                cleaned.append(line)
                seen.add(line.lower())
        return cleaned

    def _extract_docx_header_footer_lines(self, path: str) -> Tuple[List[str], List[str]]:
        def dedupe(lines: List[str]) -> List[str]:
            output: List[str] = []
            seen = set()
            for line in lines:
                key = line.lower()
                if key not in seen:
                    output.append(line)
                    seen.add(key)
            return output

        def extract_lines_from_parts(zf: zipfile.ZipFile, part_prefix: str) -> List[str]:
            collected: List[str] = []
            xml_names = [
                name for name in zf.namelist()
                if name.startswith("word/") and name.endswith(".xml") and Path(name).name.lower().startswith(part_prefix)
            ]
            for name in xml_names:
                try:
                    root = ET.fromstring(zf.read(name))
                except Exception:
                    continue

                line_parts: List[str] = []
                for node in root.iter():
                    tag_name = node.tag.rsplit("}", 1)[-1].lower()
                    if tag_name == "t":
                        txt = node.text or ""
                        if txt:
                            line_parts.append(txt)
                    elif tag_name == "tab":
                        line_parts.append("\t")
                    elif tag_name in {"br", "cr", "p"}:
                        value = self.clean_value("".join(line_parts))
                        if value:
                            collected.append(value)
                        line_parts = []
                value = self.clean_value("".join(line_parts))
                if value:
                    collected.append(value)
            return dedupe(collected)

        try:
            with zipfile.ZipFile(path) as zf:
                headers = extract_lines_from_parts(zf, "header")
                footers = extract_lines_from_parts(zf, "footer")
                try:
                    core_root = ET.fromstring(zf.read("docProps/core.xml"))
                    title_node = None
                    for node in core_root.iter():
                        if node.tag.rsplit("}", 1)[-1].lower() == "title" and (node.text or "").strip():
                            title_node = self.clean_value(node.text or "")
                            break
                    if title_node:
                        headers = dedupe([title_node] + headers)
                except Exception:
                    pass
                return headers, footers
        except Exception:
            return [], []

    def _extract_docx_text(self, path: str) -> str:
        doc = Document(path)
        parts: List[str] = []
        seen_lower = set()

        def append_part(value: str):
            value = (value or "").rstrip()
            if value == "":
                parts.append("")
                return
            lowered = value.strip().lower()
            if lowered and lowered not in seen_lower:
                parts.append(value)
                seen_lower.add(lowered)

        header_lines, footer_lines = self._extract_docx_header_footer_lines(path)
        for line in header_lines:
            append_part(line)
        if header_lines and (not parts or parts[-1] != ""):
            parts.append("")

        for para in doc.paragraphs:
            append_part(para.text or "")
        for table in doc.tables:
            if parts and parts[-1] != "":
                parts.append("")
            for row in table.rows:
                row_parts = [(cell.text or "").strip() for cell in row.cells]
                if any(row_parts):
                    append_part(" | ".join([p for p in row_parts if p]))

        textbox_lines = self._extract_docx_textbox_lines(path)
        if textbox_lines:
            if parts and parts[-1] != "":
                parts.append("")
            parts.append("[Text Boxes]")
            for line in textbox_lines:
                append_part(line)

        if footer_lines:
            if parts and parts[-1] != "":
                parts.append("")
            for line in footer_lines:
                append_part(line)

        text = "\n".join(parts)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip("\n")

    def _extract_doc_text_via_antiword(self, path: str) -> str:
        antiword = shutil.which("antiword")
        if not antiword:
            raise RuntimeError("antiword is not installed.")
        result = subprocess.run(
            [antiword, str(Path(path).resolve())],
            check=True,
            capture_output=True,
        )
        output = result.stdout.decode("utf-8", errors="ignore")
        output = output.replace("\r\n", "\n").replace("\r", "\n")
        output = re.sub(r"\n{3,}", "\n\n", output)
        output = re.sub(r"[ \t]{2,}", " ", output)
        return output.strip()


    def _extract_doc_text_via_word(self, path: str) -> str:
        try:
            import pythoncom  # type: ignore
            from win32com.client import DispatchEx  # type: ignore
        except Exception as exc:
            raise RuntimeError("Microsoft Word automation is not available.") from exc

        def unique_story_lines(raw_text: str) -> List[str]:
            # Word's COM API uses ``\x07`` (BEL) as a table-cell
            # terminator in ``Range.Text``. ``\r\x07`` marks the end of
            # a cell; lone ``\x07`` characters can also appear at the
            # end of cell content. Strip both before further processing
            # so the user never sees them as tofu squares in the
            # preview or as ``\u0007`` in the JSON export.
            raw_text = (raw_text or "").replace("\r\x07", "\n").replace("\x07", "")
            raw_text = raw_text.replace("\r", "\n")
            output: List[str] = []
            seen = set()
            for raw_line in raw_text.splitlines():
                line = self.clean_value(raw_line)
                if not line:
                    continue
                key = line.lower()
                if key not in seen:
                    output.append(line)
                    seen.add(key)
            return output

        pythoncom.CoInitialize()
        word = None
        doc = None
        try:
            word = DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0
            doc = word.Documents.Open(
                str(Path(path).resolve()),
                ConfirmConversions=False,
                ReadOnly=True,
                AddToRecentFiles=False,
                Visible=False,
            )

            header_lines: List[str] = []
            footer_lines: List[str] = []
            seen_header = set()
            seen_footer = set()

            for section_index in range(1, doc.Sections.Count + 1):
                section = doc.Sections(section_index)
                for hf_type in (1, 2, 3):
                    try:
                        header_text = section.Headers(hf_type).Range.Text or ""
                    except Exception:
                        header_text = ""
                    for line in unique_story_lines(header_text):
                        key = line.lower()
                        if key not in seen_header:
                            header_lines.append(line)
                            seen_header.add(key)

                    try:
                        footer_text = section.Footers(hf_type).Range.Text or ""
                    except Exception:
                        footer_text = ""
                    for line in unique_story_lines(footer_text):
                        key = line.lower()
                        if key not in seen_footer:
                            footer_lines.append(line)
                            seen_footer.add(key)

            # Strip Word's ``\x07`` cell-terminators alongside the
            # standard ``\r``-to-``\n`` normalisation. These BEL bytes
            # otherwise leak through into mapped field values for any
            # rule that pulls from a table cell — invisible as tofu
            # squares in the preview, exported as ``\u0007`` in JSON.
            content = (doc.Content.Text or "").replace("\r\x07", "\n").replace("\x07", "")
            content = content.replace("\r", "\n")
            content = re.sub(r"\n{3,}", "\n\n", content).strip()

            parts: List[str] = []
            if header_lines:
                parts.append("\n".join(header_lines))
            if content:
                parts.append(content)
            if footer_lines:
                parts.append("\n".join(footer_lines))

            return "\n\n".join(part for part in parts if part and part.strip()).strip()
        finally:
            if doc is not None:
                try:
                    doc.Close(False)
                except Exception:
                    pass
            if word is not None:
                try:
                    word.Quit()
                except Exception:
                    pass
            try:
                pythoncom.CoUninitialize()
            except Exception:
                pass

    def _extract_doc_text_via_soffice(self, path: str) -> str:
        soffice = shutil.which("soffice") or shutil.which("libreoffice")
        if not soffice:
            raise RuntimeError("LibreOffice is not installed.")

        with tempfile.TemporaryDirectory() as tmpdir:
            output_dir = Path(tmpdir)
            command = [
                soffice,
                "--headless",
                "--convert-to",
                "docx",
                "--outdir",
                str(output_dir),
                str(Path(path).resolve()),
            ]
            subprocess.run(command, check=True, capture_output=True)
            docx_candidates = list(output_dir.glob("*.docx"))
            if not docx_candidates:
                raise RuntimeError("LibreOffice did not produce a DOCX file.")
            return self._extract_docx_text(str(docx_candidates[0]))

    @staticmethod
    def _strip_html_tags(value: str) -> str:
        """Convert HTML to roughly-equivalent plain text.

        Handles the awkward shapes Outlook saves into ``.msg`` HTML
        bodies: VML/Office namespace elements, ``<style>``/``<script>``
        block contents, common HTML entities, and Windows-1252 smart
        quote bytes that occasionally leak through as ``\\x91`` etc.
        """
        if not value:
            return ""

        # Decode common Windows-1252 'smart' bytes that leak through when a
        # cp1252-encoded string was decoded as utf-8 with errors=ignore.
        cp1252_singles = {
            "\x91": "'",  # left single quote
            "\x92": "'",  # right single quote
            "\x93": '"',  # left double quote
            "\x94": '"',  # right double quote
            "\x96": "-",  # en dash
            "\x97": "-",  # em dash
            "\x85": "...",
        }
        for raw, replacement in cp1252_singles.items():
            value = value.replace(raw, replacement)

        # Drop entire <style> and <script> blocks (Outlook HTML is full of
        # these and their contents would otherwise leak through as raw CSS).
        value = re.sub(r"<style[^>]*>.*?</style\s*>", " ", value, flags=re.I | re.S)
        value = re.sub(r"<script[^>]*>.*?</script\s*>", " ", value, flags=re.I | re.S)

        # Drop HTML comments (often contain conditional Outlook markup).
        value = re.sub(r"<!--.*?-->", " ", value, flags=re.S)

        # Drop XML processing instructions and DOCTYPEs.
        value = re.sub(r"<\?xml[^>]*\?>", " ", value, flags=re.I)
        value = re.sub(r"<!doctype[^>]*>", " ", value, flags=re.I)

        # Block-level tags become newlines so paragraphs don't merge.
        value = re.sub(r"<br\s*/?>", "\n", value, flags=re.I)
        value = re.sub(r"</(p|div|tr|li|h[1-6])\s*>", "\n", value, flags=re.I)
        value = re.sub(r"</td\s*>", "\t", value, flags=re.I)

        # All other tags become a single space.
        value = re.sub(r"<[^>]+>", " ", value)

        # Decode HTML entities — both named (``&amp;``, ``&nbsp;``) and
        # numeric (``&#39;`` / ``&#x27;``).
        from html import unescape as _html_unescape
        value = _html_unescape(value)

        # Replace non-breaking spaces with regular ones.
        value = value.replace("\u00a0", " ")

        # Drop any leftover null bytes that occasionally appear in
        # Outlook-saved string properties.
        value = value.replace("\x00", "")

        # Tidy whitespace.
        value = re.sub(r"\n{3,}", "\n\n", value)
        value = re.sub(r"[ \t]{2,}", " ", value)
        return value.strip()

    def _extract_eml_text(self, path: str) -> str:
        with open(path, "rb") as fh:
            msg = BytesParser(policy=policy.default).parse(fh)

        parts: List[str] = []
        for header in ("Subject", "From", "To", "Date"):
            value = msg.get(header)
            if value:
                parts.append(f"{header}: {value}")
        if parts:
            parts.append("")

        body_parts: List[str] = []
        html_parts: List[str] = []
        if msg.is_multipart():
            for part in msg.walk():
                disposition = str(part.get_content_disposition() or "").lower()
                if disposition == "attachment":
                    continue
                ctype = part.get_content_type()
                try:
                    payload = part.get_content()
                except Exception:
                    try:
                        payload = part.get_payload(decode=True)
                        if isinstance(payload, bytes):
                            payload = payload.decode(part.get_content_charset() or "utf-8", errors="ignore")
                    except Exception:
                        payload = ""
                if not isinstance(payload, str):
                    continue
                if ctype == "text/plain":
                    body_parts.append(payload)
                elif ctype == "text/html":
                    html_parts.append(payload)
        else:
            payload = msg.get_content()
            if isinstance(payload, str):
                if msg.get_content_type() == "text/html":
                    html_parts.append(payload)
                else:
                    body_parts.append(payload)

        body = "\n\n".join(part.strip() for part in body_parts if part and part.strip())
        if not body and html_parts:
            body = "\n\n".join(self._strip_html_tags(part) for part in html_parts if part and part.strip())
        if body:
            parts.append(body.strip())
        text = "\n".join(parts)
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def _extract_msg_text(self, path: str) -> str:
        """Read a Microsoft Outlook .msg file using the pure-Python
        ``extract_msg`` library — no Outlook automation, no COM, no
        UI-blocking dialogs. Works for messages saved by Classic Outlook,
        New Outlook, and Outlook 365.

        The output is shaped to mirror ``_extract_eml_text`` so existing
        provider mapping rules behave the same against either format:

            Subject: ...
            From: ...
            To: ...
            Cc: ...           (only included when present)
            Date: ...

            <body>

            Attachments: a.pdf, b.docx     (only included when present)
        """
        if extract_msg is None:
            raise RuntimeError(
                "extract_msg is not installed. Run "
                "`pip install extract-msg` and try again."
            )

        msg = None
        try:
            msg = extract_msg.Message(path)

            def _coerce(value) -> str:
                """Turn an extract_msg attribute value into a clean string.

                ``extract_msg`` sometimes returns ``bytes`` (notably for
                ``htmlBody``), sometimes ``str``, sometimes a datetime,
                sometimes ``None``. String properties also occasionally
                contain trailing null bytes from the underlying compound
                file. This helper normalises all those into a stripped
                string.
                """
                if value is None:
                    return ""
                if isinstance(value, bytes):
                    for encoding in ("utf-8", "cp1252", "latin-1"):
                        try:
                            decoded = value.decode(encoding)
                            break
                        except UnicodeDecodeError:
                            continue
                    else:
                        decoded = value.decode("utf-8", errors="ignore")
                    value = decoded
                else:
                    value = str(value)
                # Strip null bytes that frequently terminate string
                # properties saved by Outlook.
                value = value.replace("\x00", "")
                return value.strip()

            def _safe_attr(attr: str) -> str:
                try:
                    return _coerce(getattr(msg, attr, None))
                except Exception:
                    return ""

            parts: List[str] = []
            for label, attr in (
                ("Subject", "subject"),
                ("From", "sender"),
                ("To", "to"),
                ("Cc", "cc"),
                ("Date", "date"),
            ):
                value = _safe_attr(attr)
                if value:
                    parts.append(f"{label}: {value}")
            if parts:
                parts.append("")

            def _looks_like_html(text: str) -> bool:
                """Some Outlook messages store HTML in the plain ``body``
                property too. Detect that so we strip tags rather than
                emitting raw markup."""
                if not text:
                    return False
                head = text.lstrip()[:200].lower()
                return ("<html" in head or "<body" in head or "<!doctype" in head
                        or "<o:p" in head or "<v:" in head)

            # Body: prefer plain text. Some Classic Outlook messages have
            # only an HTML body or only an RTF body, so fall back as needed.
            body = _safe_attr("body")
            if body and _looks_like_html(body):
                body = self._strip_html_tags(body)
            if not body:
                html = _safe_attr("htmlBody")
                if html:
                    body = self._strip_html_tags(html)
            if not body:
                # rtfBody is bytes containing RTF markup. extract_msg
                # de-encapsulates compressed RTF automatically. Strip RTF
                # control words as a best-effort fallback.
                try:
                    rtf_bytes = getattr(msg, "rtfBody", None)
                except Exception:
                    rtf_bytes = None
                if rtf_bytes:
                    if isinstance(rtf_bytes, bytes):
                        try:
                            rtf_text = rtf_bytes.decode("utf-8", errors="ignore")
                        except Exception:
                            rtf_text = ""
                    else:
                        rtf_text = str(rtf_bytes)
                    body = self._strip_rtf_markup(rtf_text)

            if body:
                parts.append(body.strip())

            # Attachment names (informational — matches what .eml users see
            # if they look at headers; helps the source preview show that
            # extra files were referenced).
            attachment_names: List[str] = []
            try:
                attachments = list(getattr(msg, "attachments", []) or [])
            except Exception:
                attachments = []
            for att in attachments:
                name = ""
                for candidate_attr in ("longFilename", "shortFilename", "displayName"):
                    try:
                        candidate = getattr(att, candidate_attr, None)
                    except Exception:
                        candidate = None
                    if candidate is None:
                        continue
                    name = _coerce(candidate)
                    if name:
                        break
                if name:
                    attachment_names.append(name)
            if attachment_names:
                parts.append("")
                parts.append("Attachments: " + ", ".join(attachment_names))

            text = "\n".join(parts)
            text = text.replace("\r\n", "\n").replace("\r", "\n")
            text = re.sub(r"\n{3,}", "\n\n", text)
            return text.strip()
        finally:
            if msg is not None:
                try:
                    msg.close()
                except Exception:
                    pass

    @staticmethod
    def _strip_rtf_markup(rtf: str) -> str:
        """Best-effort RTF -> plain text. Drops control words, group
        delimiters, and common escape sequences. Used only when a .msg
        message has no plain or HTML body — extremely rare in practice.
        """
        if not rtf:
            return ""
        # Drop binary blobs (\bin)
        text = re.sub(r"\\bin\d+\s+\S*", " ", rtf)
        # Convert hex escapes \'XX
        text = re.sub(r"\\'([0-9A-Fa-f]{2})",
                      lambda m: bytes([int(m.group(1), 16)]).decode("cp1252", errors="ignore"),
                      text)
        # Drop control words
        text = re.sub(r"\\[a-zA-Z]+-?\d*\s?", " ", text)
        # Drop control symbols
        text = re.sub(r"\\[^a-zA-Z]", "", text)
        # Drop braces
        text = text.replace("{", " ").replace("}", " ")
        # Collapse whitespace
        text = re.sub(r"[ \t]{2,}", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    @staticmethod
    def _decode_pypdf_unicode_sequences(value: str) -> str:
        if not value:
            return ""
        return re.sub(r"/uni([0-9A-Fa-f]{4})", lambda m: chr(int(m.group(1), 16)), value)

    def _extract_pdf_text_via_pymupdf(self, path: str) -> Tuple[str, List[str]]:
        """Extract PDF text using PyMuPDF in *block* mode.

        Block mode preserves the spatial grouping of text on the page: each
        visually distinct text region (e.g. a single cell in a vehicle data
        table) becomes its own block. Compared to the simpler ``"text"`` mode
        with ``sort=True``, this prevents adjacent columns from being merged
        onto the same output line — which previously caused label-based
        mapping rules to grab content from neighbouring cells (e.g.
        ``Reg No: JR07CVR Registered: May 2016 Type: Estate ...``).

        Each block's lines are joined with single newlines; blocks are
        separated from each other by a blank line so the existing label
        extraction logic still treats each block as an independent unit.
        Image blocks (``block_type == 1``) are skipped.

        OCR fallback (V54): if the document yields exactly zero text
        characters AND every page contains exactly one image, the pages
        are rendered to bitmaps and run through Tesseract OCR. This
        catches scanned-letter PDFs (single full-page image, no text
        layer) without false-triggering on photo-dump documents (which
        have multiple images per page, or some text captions, or both).
        """
        if fitz is None:
            raise RuntimeError("PyMuPDF is not installed.")
        notes: List[str] = []
        text_parts: List[str] = []
        per_page_image_counts: List[int] = []
        doc = fitz.open(path)
        try:
            for i, page in enumerate(doc, start=1):
                # Track image count per page for the OCR-fallback decision
                # below.
                try:
                    per_page_image_counts.append(len(page.get_images() or []))
                except Exception:
                    per_page_image_counts.append(0)

                blocks = page.get_text("blocks", sort=True) or []
                block_texts: List[str] = []
                for block in blocks:
                    # Each block is a tuple:
                    # (x0, y0, x1, y1, text, block_no, block_type)
                    if len(block) < 7:
                        continue
                    if block[6] != 0:
                        # Skip non-text blocks (e.g. image blocks).
                        continue
                    raw = (block[4] or "").replace("\r", "\n")
                    raw = re.sub(r"\n{2,}", "\n", raw)
                    block_lines = [line.rstrip() for line in raw.splitlines() if line.strip()]
                    if block_lines:
                        block_texts.append("\n".join(block_lines))

                if block_texts:
                    page_text = "\n\n".join(block_texts).strip()
                    if page_text:
                        text_parts.append(page_text)
                        continue

                # If block mode produced nothing usable on this page, fall
                # back to the simpler full-page text mode so we still surface
                # something for the user.
                fallback_text = page.get_text("text", sort=True) or ""
                fallback_text = fallback_text.replace("\r", "\n")
                fallback_text = re.sub(r"\n{3,}", "\n\n", fallback_text).strip()
                if fallback_text:
                    text_parts.append(fallback_text)
                else:
                    notes.append(f"Page {i} had no selectable text.")

            combined = "\n\n".join(text_parts).strip()

            # OCR fallback: only fires when the document has produced
            # *zero* text overall AND every page contains exactly one
            # image AND the document is at most ``OCR_PAGE_LIMIT`` pages.
            # All three conditions must hold so we don't waste time
            # OCR-ing photo-dump documents (which often happen to also
            # have one image per page, but are typically much longer
            # than a real scanned letter — 3+ pages versus 1-2).
            should_try_ocr = (
                not combined
                and 0 < len(per_page_image_counts) <= OCR_PAGE_LIMIT
                and all(count == 1 for count in per_page_image_counts)
                and pytesseract is not None
                and _PILImage is not None
            )
            if should_try_ocr:
                ocr_pages: List[str] = []
                ocr_failed = False
                for i, page in enumerate(doc, start=1):
                    try:
                        # Render at 300 DPI for solid OCR accuracy on
                        # typical A4 scans. ``Matrix(300/72, 300/72)``
                        # scales the default 72-DPI render up.
                        pix = page.get_pixmap(matrix=fitz.Matrix(300 / 72, 300 / 72))
                        img = _PILImage.open(io.BytesIO(pix.tobytes("png")))
                        page_ocr = pytesseract.image_to_string(img, lang="eng") or ""
                        page_ocr = page_ocr.replace("\r", "\n")
                        page_ocr = re.sub(r"\n{3,}", "\n\n", page_ocr).strip()
                        if page_ocr:
                            ocr_pages.append(page_ocr)
                    except Exception:
                        ocr_failed = True
                        break

                if ocr_pages and not ocr_failed:
                    combined = "\n\n".join(ocr_pages).strip()
                    notes.append("Read PDF using OCR fallback (no text layer).")
                    # Drop the per-page "no selectable text" notes that
                    # were collected before OCR rescued the document.
                    notes = [n for n in notes if "no selectable text" not in n]
        finally:
            doc.close()
        return combined, notes

    def extract_text(self, path: str) -> Tuple[str, List[str]]:
        notes: List[str] = []
        ext = Path(path).suffix.lower()

        if ext == ".pdf":
            if fitz is not None:
                try:
                    text, fitz_notes = self._extract_pdf_text_via_pymupdf(path)
                    notes.extend(fitz_notes)
                    notes.append("Read PDF using PyMuPDF.")
                    if text.strip():
                        return text, notes
                except Exception:
                    pass

            text_parts: List[str] = []
            try:
                reader = PdfReader(path)
                for i, page in enumerate(reader.pages, start=1):
                    page_text = self._decode_pypdf_unicode_sequences(page.extract_text() or "")
                    if page_text.strip():
                        text_parts.append(page_text)
                    else:
                        notes.append(f"Page {i} had no selectable text.")
                text = "\n\n".join(text_parts)
                if not text.strip():
                    notes.append("No selectable text found. This looks like a scanned PDF and may need OCR later.")
                else:
                    notes.append("Read PDF using pypdf.")
                return text, notes
            except Exception as exc:
                raise RuntimeError(f"Could not read PDF: {exc}") from exc

        if ext == ".docx":
            try:
                return self._extract_docx_text(path), notes
            except Exception as exc:
                raise RuntimeError(f"Could not read DOCX: {exc}") from exc

        if ext == ".doc":
            for method_name, method in (
                ("Microsoft Word", self._extract_doc_text_via_word),
                ("LibreOffice", self._extract_doc_text_via_soffice),
                ("antiword", self._extract_doc_text_via_antiword),
            ):
                try:
                    text = method(path)
                    if method_name == "antiword":
                        notes.append("Read DOC using antiword.")
                    elif method_name == "Microsoft Word":
                        notes.append("Read DOC using Microsoft Word automation.")
                    else:
                        notes.append("Read DOC using LibreOffice conversion.")
                    return text, notes
                except Exception:
                    continue
            raise RuntimeError("Could not read DOC. antiword, Microsoft Word, or LibreOffice is required for legacy .DOC files.")

        if ext == ".eml":
            try:
                return self._extract_eml_text(path), notes
            except Exception as exc:
                raise RuntimeError(f"Could not read EML: {exc}") from exc

        if ext == ".msg":
            try:
                text = self._extract_msg_text(path)
                notes.append("Read MSG using extract_msg.")
                return text, notes
            except Exception as exc:
                raise RuntimeError(f"Could not read MSG: {exc}") from exc

        raise RuntimeError("Unsupported file type. Please use PDF, DOCX, DOC, EML, or MSG.")

    def detect_provider(self, text: str) -> ProviderMatch:
        """Pick the best matching provider for ``text``.

        Detection rule: every phrase in a provider's ``detect_phrases``
        list must be present in the document. A provider with a longer
        / more specific fingerprint (more phrases) naturally wins over
        a shorter one when both match, because each matching phrase
        contributes to the score and the tiebreakers prefer the
        higher-match-count provider. This is what lets two providers
        share a base phrase like ``fairwaylegal`` while one of them
        adds extra phrases (e.g. ``Inspection Location:``) to handle
        a more specific document shape.

        Providers with no configured phrases are skipped.
        """
        lower_text = text.lower()
        normalized_text = self.normalize_search_text(text)
        best = ProviderMatch("Unknown / Unmapped", 0, {"field_rules": {}})
        best_match_count = 0
        best_longest = 0

        for provider in self.config.get("providers", []):
            phrases = [str(phrase).strip() for phrase in provider.get("detect_phrases", []) if str(phrase).strip()]
            if not phrases:
                continue

            score = 0
            match_count = 0
            longest_match = 0
            all_matched = True

            for phrase in phrases:
                raw_phrase = phrase.lower().strip()
                normalized_phrase = self.normalize_search_text(phrase)
                matched = False

                if raw_phrase and raw_phrase in lower_text:
                    matched = True
                elif normalized_phrase and normalized_phrase in normalized_text:
                    matched = True

                if matched:
                    match_count += 1
                    longest_match = max(longest_match, len(normalized_phrase or raw_phrase))
                    score += max(10, len(normalized_phrase or raw_phrase))
                else:
                    all_matched = False
                    break

            if not all_matched:
                continue

            if (
                score > best.score
                or (score == best.score and match_count > best_match_count)
                or (score == best.score and match_count == best_match_count and longest_match > best_longest)
            ):
                best = ProviderMatch(provider.get("name", "Unknown / Unmapped"), score, provider)
                best_match_count = match_count
                best_longest = longest_match

        return best

    @staticmethod
    def clean_value(value: str) -> str:
        value = value.replace("\r", " ").replace("\t", " ")
        value = re.sub(r"[ ]{2,}", " ", value)
        value = re.sub(r"\n{3,}", "\n\n", value)
        # Strip surrounding whitespace and trailing/leading colons.
        # Hyphens are deliberately *not* stripped — users sometimes
        # legitimately want a single ``-`` (or a string ending in ``-``)
        # as a label in two-label / single-label mappings.
        return value.strip(" :\n")

    @staticmethod
    def normalize_search_text(value: str) -> str:
        value = value.lower().replace("\r", "\n").replace("\t", " ")
        value = re.sub(r"[^\w\n ]+", " ", value)
        value = re.sub(r"\s+", " ", value)
        return value.strip()

    def lines(self, text: str) -> List[str]:
        return [self.clean_value(line) for line in text.splitlines()]

    def raw_lines(self, text: str) -> List[str]:
        return [line.replace("\r", "").replace("\t", " ").strip() for line in text.splitlines()]

    @staticmethod
    def line_looks_like_new_label(line: str) -> bool:
        if not line:
            return False
        if re.match(r"^[A-Za-z][A-Za-z /&()'’\-]{1,60}:$", line):
            return True
        if len(line) < 60 and line.count(":") == 1:
            left, _ = line.split(":", 1)
            return 0 < len(left.strip()) < 35
        return False

    @staticmethod
    def looks_like_section_heading(line: str) -> bool:
        stripped = line.strip()
        if not stripped:
            return False
        if len(stripped) > 80:
            return False
        return bool(re.match(r"^[A-Z][A-Z0-9 '&()\-/]{6,}$", stripped))

    @staticmethod
    def extract_uk_postcode(line: str) -> str:
        if not line:
            return ""
        match = re.search(r"\b([A-Z]{1,2}\d[A-Z\d]?\s*\d[ABD-HJLNP-UW-Z]{2})\b", line, re.IGNORECASE)
        if not match:
            return ""
        compact = re.sub(r"\s+", "", match.group(1).upper())
        if len(compact) < 5:
            return ""
        return f"{compact[:-3]} {compact[-3:]}"

    def text_blocks(self, text: str) -> List[List[str]]:
        blocks = []
        current = []
        for raw in self.raw_lines(text):
            cleaned = self.clean_value(raw)
            if cleaned:
                current.append(cleaned)
            elif current:
                blocks.append(current)
                current = []
        if current:
            blocks.append(current)
        return blocks

    def extract_multiline_after_labels(self, text: str, labels: List[str], max_lines: int = 6) -> str:
        if not labels:
            return ""
        labels_clean = [self.clean_value(label) for label in labels if self.clean_value(label)]
        if not labels_clean:
            return ""

        lines = self.raw_lines(text)
        for i, line in enumerate(lines):
            norm_line = self.clean_value(line).lower()
            matched = False
            for label in labels_clean:
                norm_label = label.lower()
                if norm_line == norm_label or norm_line.rstrip(":") == norm_label:
                    matched = True
                elif norm_line.startswith(norm_label + ":"):
                    remainder = self.clean_value(line.split(":", 1)[1])
                    if remainder:
                        return remainder
                    matched = True
                elif norm_line.startswith(norm_label) and len(norm_line) <= len(norm_label) + 2:
                    matched = True
                if matched:
                    break

            if matched:
                collected: List[str] = []
                blank_run = 0
                for next_line in lines[i + 1:]:
                    cleaned = self.clean_value(next_line)
                    if not cleaned:
                        blank_run += 1
                        if collected and blank_run >= 2:
                            break
                        continue
                    blank_run = 0
                    if collected and (self.line_looks_like_new_label(cleaned) or self.looks_like_section_heading(cleaned)):
                        break
                    if collected and re.search(r"^(Dear Sirs|Yours faithfully|Please arrange|I act on behalf|Finally having regard)", cleaned, re.I):
                        break
                    collected.append(cleaned)
                    if self.extract_uk_postcode(cleaned):
                        break
                    if len(collected) >= max_lines:
                        break
                if collected:
                    return "\n".join(collected)
        return ""

    def extract_after_label(self, text: str, labels: List[str], multiline: bool = False) -> str:
        if multiline:
            value = self.extract_multiline_after_labels(text, labels)
            if value:
                return value

        if not labels:
            return ""

        labels_clean = [self.clean_value(label) for label in labels if self.clean_value(label)]
        if not labels_clean:
            return ""

        lines = self.lines(text)
        label_union = "|".join(re.escape(label) for label in labels_clean)

        same_line_patterns = [
            re.compile(rf"(?im)^\s*(?:{label_union})\s*[:#\-|]?\s*(.+?)\s*$"),
            re.compile(rf"(?im)\b(?:{label_union})\b\s*[:#\-|]?\s*(.+)"),
        ]
        lower_labels = {label.lower() for label in labels_clean}

        for regex in same_line_patterns:
            for match in regex.finditer(text):
                value = self.clean_value(match.group(1))
                if value and value.lower() not in lower_labels:
                    return value

        for i, line in enumerate(lines):
            if not line:
                continue
            matched = False
            for label in labels_clean:
                norm_label = label.lower()
                norm_line = line.lower()
                if norm_line == norm_label or norm_line.rstrip(":") == norm_label:
                    matched = True
                elif norm_line.startswith(norm_label + ":") or norm_line.startswith(norm_label + " -"):
                    remainder = self.clean_value(line[len(label):])
                    if remainder:
                        return remainder
                    matched = True
                elif norm_line.startswith(norm_label) and len(line) <= len(label) + 2:
                    matched = True
                if matched:
                    break

            if matched:
                collected: List[str] = []
                for next_line in lines[i + 1:]:
                    if not next_line:
                        if collected:
                            break
                        continue
                    if self.line_looks_like_new_label(next_line) and collected:
                        break
                    collected.append(next_line)
                    if not multiline:
                        break
                    if len(" ".join(collected)) > 250:
                        break
                if collected:
                    return self.clean_value("\n".join(collected) if multiline else collected[0])
        return ""

    def first_regex_match(self, text: str, patterns: List[str]) -> str:
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                return self.clean_value(match.group(1))
        return ""

    def extract_reference_fallback(self, text: str) -> str:
        return self.first_regex_match(
            text,
            [
                r"(?im)^\s*(?:our\s+reference|our\s+ref|reference|ref(?:erence)?)\s*[:#\-]?\s*([^\n]+)",
                r"(?im)^\s*([A-Z0-9./_-]{6,})\s*$",
            ],
        )

    def extract_vrm_fallback(self, text: str) -> str:
        lines = self.lines(text)
        for line in lines[:120]:
            if re.search(r"(registration|vehicle registration|vrm|reg)", line, re.I):
                match = re.search(r"\b([A-Z]{2}\d{2}[A-Z]{3}|[A-Z]\d{1,3}[A-Z]{3}|[A-Z]{3}\d{1,3}[A-Z]|\d{1,4}[A-Z]{1,3})\b", line, re.I)
                if match:
                    return match.group(1).upper()
        match = re.search(r"\b([A-Z]{2}\d{2}[A-Z]{3}|[A-Z]\d{1,3}[A-Z]{3}|[A-Z]{3}\d{1,3}[A-Z]|\d{1,4}[A-Z]{1,3})\b", text, re.I)
        return match.group(1).upper() if match else ""

    def extract_name_fallback(self, text: str) -> str:
        blocks = self.text_blocks(text)
        titled_name = re.compile(r"^(Mr|Mrs|Miss|Ms|Dr|Mx|Prof)\.?\s+[A-Z][A-Za-z'’\-]+(?:\s+[A-Z][A-Za-z'’\-]+){0,4}$")
        for block in blocks[:20]:
            for line in block:
                if titled_name.match(line):
                    return line
        return self.first_regex_match(
            text,
            [
                r"(?im)^\s*(?:our\s+client|client|claimant)\s*[:#\-]?\s*([^\n]+)",
            ],
        )

    def extract_vehicle_model_fallback(self, text: str) -> str:
        value = self.first_regex_match(
            text,
            [
                r"(?im)^\s*(?:client\s+vehicle\s+model|vehicle\s+model|model)\s*[:#\-]?\s*([^\n]+)",
                r"(?im)^\s*(?:our\s+client.?s\s+vehicle|client.?s\s+vehicle|vehicle details?)\s*[:#\-]?\s*([^\n]+)",
            ],
        )
        if value:
            value = re.sub(r"\b([A-Z]{2}\d{2}[A-Z]{3})\b", "", value, flags=re.I).strip()
        return value

    def extract_date_near_keywords(self, text: str, keywords: List[str]) -> str:
        keyword_union = "|".join(re.escape(k) for k in keywords if k)
        pattern = re.compile(
            rf"(?im)(?:{keyword_union})\s*[:#\-]?\s*(\d{{1,2}}(?:st|nd|rd|th)?\s+[A-Za-z]+\s+\d{{4}}|\d{{1,2}}[/-]\d{{1,2}}[/-]\d{{2,4}})",
            re.IGNORECASE,
        )
        match = pattern.search(text)
        return self.clean_value(match.group(1)) if match else ""

    def extract_address_fallback(self, text: str) -> str:
        value = self.extract_multiline_after_labels(text, ["Address", "Inspection Address", "Inspection Location"], max_lines=7)
        if value:
            return value

        lines = self.raw_lines(text)
        anchors = [r"available at\s*:?$", r"inspection (?:will be|to be)?\s*at\s*:?$", r"vehicle is located at\s*:?$"]
        for i, line in enumerate(lines):
            cleaned = self.clean_value(line)
            if any(re.search(pat, cleaned, re.I) for pat in anchors):
                collected = []
                blank_run = 0
                for next_line in lines[i + 1:]:
                    cleaned2 = self.clean_value(next_line)
                    if not cleaned2:
                        blank_run += 1
                        if collected and blank_run >= 2:
                            break
                        continue
                    blank_run = 0
                    if collected and self.line_looks_like_new_label(cleaned2):
                        break
                    if collected and re.search(r"^(Mobile Tel|Home Tel|Work Tel|Tel|Email)", cleaned2, re.I):
                        break
                    collected.append(cleaned2)
                    if self.extract_uk_postcode(cleaned2):
                        break
                    if len(collected) >= 7:
                        break
                if collected:
                    return "\n".join(collected)
        return ""

    def extract_letterhead_date(self, text: str) -> str:

        long_date = re.compile(
            r"\b(\d{1,2}(?:st|nd|rd|th)?\s+[A-Za-z]+\s+\d{4}|\d{1,2}[/-]\d{1,2}[/-]\d{2,4})\b",
            re.IGNORECASE,
        )
        candidates = []
        for line in self.lines(text)[:25]:
            if not line:
                continue
            if ":" in line and len(line.split(":", 1)[0].strip()) < 25:
                continue
            match = long_date.search(line)
            if match:
                candidates.append(self.clean_value(match.group(1)))
        return candidates[0] if candidates else ""

    def parse_two_label_config(self, config_value: str) -> Tuple[str, str]:
        """Split a ``"start || end"`` user-config string into its two halves.

        The user's literal characters are preserved — only surrounding
        whitespace is trimmed. Specifically, leading/trailing colons and
        hyphens are NOT stripped here, because users sometimes
        legitimately want a single ``:`` or ``-`` (or a label ending in
        one of those characters) as a label. The downstream matching
        code does its own normalisation if it needs to.
        """
        raw = (config_value or "").strip()
        if "||" in raw:
            start, end = raw.split("||", 1)
            return start.strip(), end.strip()
        parts = [part.strip() for part in raw.splitlines() if part.strip()]
        if len(parts) >= 2:
            return parts[0], parts[1]
        return "", ""

    def extract_between_labels(self, text: str, start_label: str, end_label: str) -> str:
        if not start_label or not end_label:
            return ""
        pattern = re.compile(
            rf"(?is){re.escape(start_label)}\s*:?\s*(.*?)\s*(?={re.escape(end_label)})"
        )
        match = pattern.search(text)
        if match:
            value = self.clean_value(match.group(1))
            return value

        lines = text.splitlines()
        capture = False
        collected: List[str] = []
        for raw_line in lines:
            line = raw_line.rstrip()
            cleaned = self.clean_value(line)
            lower = cleaned.lower()
            if not capture:
                if lower.startswith(start_label.lower()):
                    capture = True
                    remainder = self.clean_value(re.sub(rf"(?i)^{re.escape(start_label)}\s*:?", "", cleaned))
                    if remainder:
                        collected.append(remainder)
                    continue
            else:
                if lower.startswith(end_label.lower()):
                    break
                collected.append(cleaned)
        return self.clean_value("\n".join([part for part in collected if part]))

    def extract_by_fixed_position(self, text: str, config_value: str) -> str:
        raw = (config_value or "").strip()
        match = re.match(r"^(\d+)\s*(?:-|:)\s*(\d+)$", raw) or re.match(r"^(\d+)$", raw)
        if not match:
            return ""
        start = int(match.group(1))
        end = int(match.group(2)) if match.lastindex and match.lastindex >= 2 and match.group(2) else start
        if start <= 0 or end < start:
            return ""
        lines = text.splitlines()
        if start > len(lines):
            return ""
        selected = [self.clean_value(line) for line in lines[start - 1:end]]
        selected = [line for line in selected if line]
        return self.clean_value("\n".join(selected))

    def extract_by_fixed_position_label(self, text: str, config_value: str) -> str:
        """Look at one specific line of the source text and extract the
        portion of that line after the user's label.

        Config format mirrors the Two Labels approach: ``"<line> || <label>"``
        where ``<line>`` is a 1-based line number (matching the source
        preview's line numbering) and ``<label>`` is the literal text to
        find inside that line. The returned value is everything after
        the label on that line, trimmed.

        If the line number is out of range, the line doesn't contain the
        label, or the config is malformed, returns "".

        Useful when a label like ``-`` would be too ambiguous to use as
        a Single Label across the whole document, but is unambiguous
        once narrowed to a specific line — e.g.::

            Line 4: Toyota Prius - BR19 SRX

        Configured as ``4 || -`` returns ``BR19 SRX``.
        """
        raw = (config_value or "").strip()
        if not raw:
            return ""
        line_part, label = self.parse_two_label_config(raw)
        if not line_part or not label:
            return ""
        try:
            line_no = int(line_part)
        except (TypeError, ValueError):
            return ""
        if line_no <= 0:
            return ""
        lines = text.splitlines()
        if line_no > len(lines):
            return ""
        target_line = lines[line_no - 1]
        # Case-insensitive search for the label within the target line.
        idx = target_line.lower().find(label.lower())
        if idx < 0:
            return ""
        after = target_line[idx + len(label):]
        return self.clean_value(after)

    def extract_by_single_label_offset(self, text: str, config_value: str) -> str:
        """Find a label anywhere in the document, then return the line
        ``offset`` *non-blank* lines above or below it.

        Config format: ``"<label> || <offset>"`` where ``<offset>`` is
        a signed integer like ``-2``, ``+1``, or ``0``. The sign is
        required (the strict ``+/-`` form makes the intent explicit).

        Blank lines are skipped when counting the offset. So
        ``label || +1`` returns the first non-blank line below the
        label, ``+2`` the second, etc. ``+0`` and ``-0`` return the
        anchor (label) line itself.

        Useful for OCR'd documents where a value sits on a predictable
        line near a recognisable label, but where Fixed Position alone
        won't work because OCR line numbering varies between scans —
        and especially useful when OCR introduces stray blank lines
        between rows that would otherwise throw off a strict line
        offset.

        Out-of-bounds offsets clamp to the last/first non-blank line in
        that direction. The rule only returns "" if the label can't be
        found at all, the offset isn't a signed integer, or the config
        is malformed.
        """
        raw = (config_value or "").strip()
        if not raw:
            return ""
        label, offset_part = self.parse_two_label_config(raw)
        if not label or not offset_part:
            return ""
        # Strict +/- form: offset must start with + or -, then digits.
        m = re.fullmatch(r"\s*([+\-])\s*(\d+)\s*", offset_part)
        if not m:
            return ""
        sign = m.group(1)
        magnitude = int(m.group(2))
        offset = magnitude if sign == "+" else -magnitude

        lines = text.splitlines()
        if not lines:
            return ""
        # Find the FIRST line containing the label, case-insensitive.
        needle = label.lower()
        anchor: Optional[int] = None
        for i, line in enumerate(lines):
            if needle in line.lower():
                anchor = i
                break
        if anchor is None:
            return ""

        if offset == 0:
            return self.clean_value(lines[anchor])

        # Walk through ``lines`` skipping blank ones, counting only
        # non-blank lines toward the requested offset. ``last_seen``
        # tracks the most recent non-blank line we've stepped onto so
        # we can clamp gracefully if we run out of document.
        step = 1 if offset > 0 else -1
        steps_remaining = abs(offset)
        i = anchor + step
        last_seen = anchor
        while 0 <= i < len(lines) and steps_remaining > 0:
            if lines[i].strip():
                last_seen = i
                steps_remaining -= 1
                if steps_remaining == 0:
                    break
            i += step
        return self.clean_value(lines[last_seen])

    def extract_by_email_date(self, text: str, config_value: str) -> str:
        """Find a label, then look on the same line for a YYYY-MM-DD
        date and return it formatted as DD/MM/YYYY.

        Designed for ``.msg`` email headers like::

            Date: 2026-05-05 16:57:31+01:00

        where the date is on the label line itself in ISO form. The
        config takes a single label string (e.g. ``Date:``); the
        method finds the first line containing that label, scans the
        portion after the label for the first ``YYYY-MM-DD`` match,
        and returns it as ``DD/MM/YYYY``.

        Looks at the *label line only* — does not fall through to the
        next line. This is deliberate: ``Date:`` labels also appear
        in instruction-letter bodies on their own line, and we don't
        want to accidentally match document body content as if it
        were an email header.

        Returns "" if the label isn't found, no YYYY-MM-DD pattern
        is on the label line, or the matched date isn't a real date.
        """
        raw_label = (config_value or "").strip()
        if not raw_label:
            return ""
        needle = raw_label.lower()
        # Match YYYY-MM-DD where YYYY is 4 digits, MM is 1-2, DD is 1-2.
        # We post-validate via strptime so 2026-13-45 etc. fail.
        date_re = re.compile(r"\b(\d{4}-\d{1,2}-\d{1,2})\b")
        for line in text.splitlines():
            idx = line.lower().find(needle)
            if idx < 0:
                continue
            # Search only the portion of the line after the label.
            tail = line[idx + len(raw_label):]
            match = date_re.search(tail)
            if not match:
                # Label found, but no date on this line — keep
                # scanning in case the label appears again later.
                continue
            iso = match.group(1)
            try:
                dt = datetime.strptime(iso, "%Y-%m-%d")
            except ValueError:
                continue
            return dt.strftime("%d/%m/%Y")
        return ""


    def config_tokens(self, config_value: str) -> List[str]:
        return [self.clean_value(part) for part in (config_value or "").split(",") if self.clean_value(part)]

    def extract_by_rule(self, text: str, field_name: str, rule: dict) -> str:
        rule = rule or {}
        method = rule.get("method") or DEFAULT_METHOD_BY_FIELD.get(field_name, "single_label")
        config_value = rule.get("config", "")
        tokens = self.config_tokens(config_value)

        if method == "single_label":
            return self.extract_after_label(text, tokens, multiline=False)
        if method == "two_labels":
            start_label, end_label = self.parse_two_label_config(config_value)
            value = self.extract_between_labels(text, start_label, end_label)
            if field_name.endswith("_date") and value:
                date_match = re.search(r"\b(\d{1,2}(?:st|nd|rd|th)?\s+[A-Za-z]+\s+\d{4}|\d{1,2}[/-]\d{1,2}[/-]\d{2,4})\b", value, re.IGNORECASE)
                if date_match:
                    return self.clean_value(date_match.group(1))
            return value
        if method == "fixed_position":
            return self.extract_by_fixed_position(text, config_value)
        if method == "fixed_position_label":
            return self.extract_by_fixed_position_label(text, config_value)
        if method == "single_label_offset":
            return self.extract_by_single_label_offset(text, config_value)
        if method == "email_date":
            return self.extract_by_email_date(text, config_value)
        if method == "manual_input":
            if config_value.strip().lower() == "{today}":
                return datetime.now().strftime("%d/%m/%Y")
            return config_value.strip()

        # Backwards-compatible handling for older provider files.
        if method == "labels":
            return self.extract_after_label(text, tokens, multiline=False)
        if method == "multiline_labels":
            return self.extract_after_label(text, tokens, multiline=True)
        if method == "letterhead_date":
            return self.extract_letterhead_date(text)
        if method == "date_near_keywords":
            return self.extract_date_near_keywords(text, tokens)
        if method == "reference_fallback":
            return self.extract_reference_fallback(text)
        if method == "vrm_fallback":
            return self.extract_vrm_fallback(text)
        if method == "claimant_name_fallback":
            return self.extract_name_fallback(text)
        if method == "vehicle_model_fallback":
            return self.extract_vehicle_model_fallback(text)
        if method == "address_fallback":
            return self.extract_address_fallback(text)
        if method == "fixed_value":
            return config_value.strip()
        if method == "current_date":
            return datetime.now().strftime("%d/%m/%Y")
        if method == "blank":
            return ""
        return self.extract_after_label(text, tokens, multiline=False)


    @staticmethod
    def normalize_yes_no(value: str) -> str:
        lowered = (value or "").strip().lower()
        if not lowered:
            return ""
        if lowered in {"yes", "y", "true", "1"}:
            return "Yes"
        if lowered in {"no", "n", "false", "0"}:
            return "No"
        return value.strip()

    def resolve_presence_check(self, text: str, rule: dict, positive_value: str, negative_value: str) -> str:
        """Generic Yes/No-style field resolver.

        The user enters one or more comma-separated tokens describing the
        *positive* scenario. If any token appears anywhere in the document
        text (case-insensitive), the field is set to ``positive_value``;
        otherwise to ``negative_value``. A blank config returns "".

        There is no fallback chain and no method-dependent behaviour — this
        deliberately gives the user full control. Used for VAT Status
        (Yes/No) and Mileage Unit (Miles/Km).
        """
        rule = rule or {}
        config_value = str(rule.get("config", "") or "").strip()
        if not config_value:
            return ""

        tokens = [self.clean_value(part) for part in config_value.split(",") if self.clean_value(part)]
        if not tokens:
            return negative_value

        haystack = (text or "").lower()
        for token in tokens:
            needle = token.lower()
            if needle and needle in haystack:
                return positive_value
        return negative_value

    def extract_fields(self, text: str, provider: ProviderMatch) -> Tuple[Dict[str, str], List[str]]:
        notes: List[str] = []
        values = {key: "" for key, _ in DEFAULT_FIELDS}

        provider_config = provider.config or {}
        field_rules = (provider_config.get("field_rules") or {})
        use_current_date_for_inspection_date = bool(provider_config.get("use_current_date_for_inspection_date"))
        force_postcode_for_inspection_address = bool(provider_config.get("force_postcode_for_inspection_address"))

        # Work Provider is a manual-input-only field controlled by the
        # preset itself. The user types it once when creating the preset
        # (e.g. "FW") and that value is what populates the Detected
        # Fields panel and the JSON export — independent of the preset
        # name (which can be more descriptive, e.g. "FW (Garage)" /
        # "FW (Solicitor)"). If the preset has no work_provider rule
        # set, the field stays blank and the user knows to update the
        # preset.
        if provider.name != "Unknown / Unmapped":
            wp_rule = field_rules.get("work_provider") or {}
            wp_value = str((wp_rule or {}).get("config", "") or "").strip()
            values["work_provider"] = wp_value

        for field_name in NON_PROVIDER_FIELDS:
            rule = field_rules.get(field_name) or {
                "method": DEFAULT_METHOD_BY_FIELD.get(field_name, "single_label"),
                "config": "",
            }
            config_value = str((rule or {}).get("config", "") or "").strip()

            # Presets now have full control. If the config is blank, do not map anything
            # for that field and do not allow any fallback extraction to kick in.
            if not config_value:
                continue

            if field_name in PRESENCE_CHECK_FIELDS:
                pc = PRESENCE_CHECK_FIELDS[field_name]
                values[field_name] = self.resolve_presence_check(
                    text, rule, pc["positive_value"], pc["negative_value"]
                )
                continue

            extracted = self.extract_by_rule(text, field_name, rule)
            if extracted:
                extracted = post_process_extracted_value(
                    field_name,
                    extracted,
                    force_postcode=force_postcode_for_inspection_address if field_name == "inspection_address" else False,
                )
                values[field_name] = extracted

        if use_current_date_for_inspection_date:
            values["inspection_date"] = datetime.now().strftime("%d/%m/%Y")

        missing_required = [FIELD_LABELS[key] for key in REQUIRED_FIELDS if not values.get(key)]
        if provider.name == "Unknown / Unmapped":
            notes.append("Provider not recognised yet.")
        if missing_required:
            notes.append("Missing required fields: " + ", ".join(sorted(missing_required)))
        return values, notes
    def extract_images_to_desktop(self, path: str, output_dir: Path, base_name: str) -> Tuple[List[Path], List[str]]:
        source = Path(path)
        ext = source.suffix.lower()
        saved: List[Path] = []
        notes: List[str] = []

        def save_bytes(stem: str, suffix: str, data: bytes) -> Path:
            out = unique_output_path(output_dir, stem, suffix)
            out.write_bytes(data)
            return out

        def extract_from_docx(docx_path: Path):
            with zipfile.ZipFile(docx_path, "r") as zf:
                media = [name for name in zf.namelist() if name.startswith("word/media/") and not name.endswith("/")]
                for idx, member in enumerate(media, start=1):
                    suffix = Path(member).suffix or ".bin"
                    data = zf.read(member)
                    saved.append(save_bytes(f"{base_name}_img_{idx}", suffix, data))

        try:
            if ext == ".pdf":
                reader = PdfReader(str(source))
                idx = 1
                for page_num, page in enumerate(reader.pages, start=1):
                    images = getattr(page, "images", []) or []
                    for img in images:
                        suffix = Path(getattr(img, "name", "")).suffix or ".bin"
                        saved.append(save_bytes(f"{base_name}_img_{page_num}_{idx}", suffix, img.data))
                        idx += 1
                if not saved:
                    notes.append("No extractable images were found in the PDF.")
            elif ext == ".docx":
                extract_from_docx(source)
                if not saved:
                    notes.append("No embedded images were found in the DOCX.")
            elif ext == ".doc":
                with tempfile.TemporaryDirectory() as tmpdir:
                    tmpdir_path = Path(tmpdir)
                    temp_docx = tmpdir_path / f"{source.stem}.docx"
                    converted = None
                    try:
                        convert_doc_to_docx_via_word(source, temp_docx)
                        converted = temp_docx
                    except Exception:
                        try:
                            converted = convert_doc_to_docx_via_soffice(source, tmpdir_path)
                        except Exception as exc:
                            notes.append(f"Could not convert DOC for image extraction: {exc}")
                    if converted and converted.exists():
                        extract_from_docx(converted)
                        if not saved:
                            notes.append("No embedded images were found in the DOC.")
            else:
                notes.append("Image extraction is only supported for PDF, DOCX, and DOC.")
        except Exception as exc:
            notes.append(f"Image extraction failed: {exc}")

        if saved:
            notes.append(f"Extracted {len(saved)} image(s) to Desktop.")
        return saved, notes

def split_preserving_nonempty_lines(text: str) -> List[str]:
    return [part.strip() for part in (text or "").splitlines() if part.strip()]


def normalise_inspection_address_value(value: str, force_postcode: bool = False) -> str:
    """Normalise the inspection address to a 6-line canonical form.

    The output is always exactly 6 lines separated by 5 newlines —
    ``"line1\\nline2\\nline3\\nline4\\nline5\\npostcode"`` — including
    when the input is empty (in which case all 6 lines are blank,
    yielding ``"\\n\\n\\n\\n\\n"``). Downstream JSON export feeds this
    to a management system that requires the 6-line shape; returning
    a bare empty string would fail the import.

    Lines 1-5 are the address body, line 6 is the postcode. Body lines
    overflow by joining surplus content into line 5 with spaces.
    """
    text = (value or "").strip()
    if not text:
        # Always emit the 6-line canonical shape, even for empty input.
        return "\n".join([""] * 6)

    def _canonicalise_postcode(postcode_text: str) -> str:
        compact = re.sub(r"\s+", "", (postcode_text or "").upper())
        if len(compact) < 5:
            return ""
        return f"{compact[:-3]} {compact[-3:]}"

    postcode_anywhere_re = re.compile(r"\b([A-Z]{1,2}\d[A-Z\d]?\s*\d[ABD-HJLNP-UW-Z]{2})\b", re.IGNORECASE)
    postcode_end_re = re.compile(r"\b([A-Z]{1,2}\d[A-Z\d]?\s*\d[ABD-HJLNP-UW-Z]{2})\b\s*$", re.IGNORECASE)

    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\s*,\s*", "\n", text)
    raw_lines = split_preserving_nonempty_lines(text)
    if not raw_lines:
        return "\n".join([""] * 6)

    postcode_line = ""
    body_lines: List[str] = []

    if len(raw_lines) == 1:
        single_line = raw_lines[0]
        if force_postcode:
            end_match = postcode_end_re.search(single_line)
            if end_match:
                postcode_line = _canonicalise_postcode(end_match.group(1))
                remainder = single_line[:end_match.start()].strip(" ,")
                body_lines = [remainder] if remainder else []
            else:
                # Postcode isn't at end-of-line — common when the address
                # ends with a phone number, e.g.
                #   "Somstar Recovery ... Birmingham B5 6JX 07462530375"
                # Find the postcode anywhere in the line and split there:
                # everything before becomes the body, the postcode goes
                # in line 6, and any trailing content (typically a phone
                # number) is dropped.
                anywhere_match = postcode_anywhere_re.search(single_line)
                if anywhere_match:
                    postcode_line = _canonicalise_postcode(anywhere_match.group(1))
                    pre = single_line[:anywhere_match.start()].strip(" ,")
                    body_lines = [pre] if pre else []
                else:
                    body_lines = raw_lines[:]
        else:
            body_lines = raw_lines[:]
    else:
        last_line = raw_lines[-1]
        any_match = postcode_anywhere_re.search(last_line)
        postcode_line = _canonicalise_postcode(any_match.group(1)) if any_match else last_line
        body_lines = raw_lines[:-1]

    if len(body_lines) >= 5:
        line1 = body_lines[0]
        line2 = body_lines[1] if len(body_lines) > 1 else ""
        line3 = body_lines[2] if len(body_lines) > 2 else ""
        line4 = body_lines[3] if len(body_lines) > 3 else ""
        overflow = [part for part in body_lines[4:] if part]
        line5 = " ".join(overflow)
        normalized = [line1, line2, line3, line4, line5, postcode_line]
    else:
        body_lines = body_lines[:5]
        normalized = body_lines + [""] * (5 - len(body_lines)) + [postcode_line]

    normalized = [part.strip() for part in normalized[:6]]
    while len(normalized) < 6:
        normalized.append("")
    return "\n".join(normalized)


def post_process_extracted_value(field_name: str, value: str, force_postcode: bool = False) -> str:
    if field_name == "vrm":
        return normalize_vrm_value(value)
    if field_name == "inspection_address":
        return normalise_inspection_address_value(value, force_postcode=force_postcode)
    if field_name == "mileage":
        return normalize_mileage_value(value)
    return value


def normalise_date_value(value: str) -> str:
    """Convert a date string into ``DD/MM/YYYY`` form.

    Tries a list of known formats in priority order. On success
    returns ``DD/MM/YYYY``; on failure (gibberish, ambiguous, or
    unsupported format) returns the original input unchanged so the
    user can spot the issue at JSON-import time and fix the source.

    Recognised formats:

    - ``DD/MM/YYYY`` (already canonical, passes through)
    - ``DD/MM/YY``  → expanded to 4-digit year via ``strptime``'s
      default 2-digit-year handling
    - ``DD-MM-YYYY`` and ``DD-MM-YY``
    - ``DD MMMM YYYY`` (e.g. ``27 April 2026``)
    - ``DD MMM YYYY``  (e.g. ``21 Apr 2026``)
    - ``MMMM DD YYYY`` (e.g. ``April 27 2026``)
    - ``MMM DD YYYY``  (e.g. ``Apr 21 2026``)
    - ``YYYY-MM-DD``   (ISO)

    Ordinal suffixes (``1st``, ``2nd``, ``3rd``, ``27th``) are
    stripped before parsing, so ``27th April 2026`` is recognised
    as ``27 April 2026``.

    Whitespace, commas, and surrounding punctuation are tolerated.

    This is called only at JSON-export time so the Detected Fields
    panel preserves whatever shape the document originally used —
    the user only sees normalised dates in the exported JSON.
    """
    raw = (value or "").strip()
    if not raw:
        return ""

    # Strip ordinal suffixes ("27th" -> "27", "9 th" -> "9") and
    # commas, normalise whitespace. The ``\s*`` between the digits
    # and suffix lets us handle some real-world documents that emit
    # the day with a space before the suffix. ``\b`` after the
    # suffix prevents false matches like "5 thousand" (the position
    # after ``th`` is the word-character ``o``, so the boundary
    # fails). Separators ``/`` and ``-`` are untouched because the
    # format strings rely on them.
    cleaned = re.sub(r"(\d+)\s*(st|nd|rd|th)\b", r"\1", raw, flags=re.IGNORECASE)
    cleaned = cleaned.replace(",", " ")
    cleaned = re.sub(r"\s+", " ", cleaned).strip()

    formats = [
        "%d/%m/%Y",
        "%d/%m/%y",
        "%d-%m-%Y",
        "%d-%m-%y",
        "%d %B %Y",
        "%d %b %Y",
        "%B %d %Y",
        "%b %d %Y",
        "%Y-%m-%d",
    ]
    for fmt in formats:
        try:
            dt = datetime.strptime(cleaned, fmt)
        except ValueError:
            continue
        return dt.strftime("%d/%m/%Y")

    # Couldn't parse — leave unchanged so the user notices the issue
    # at JSON-import time and fixes the source.
    return raw


def normalize_vrm_value(value: str) -> str:
    """Normalize extracted VRMs by removing all whitespace and uppercasing."""
    return re.sub(r"\s+", "", (value or "").strip()).upper()


def normalize_mileage_value(value: str) -> str:
    """Extract a mileage number from the text returned by the user's mapping rule.

    Walks left-to-right through the value, ignoring leading whitespace and any
    other non-digit / non-comma characters before the first digit. Once digits
    start, collects digits and commas (commas treated as in-number separators
    and stripped from the result), and stops at the first character that is
    neither a digit nor a comma. Returns "" if no digits are found.

    Example: "   28,487 Miles" -> "28487"
             "Speedo: 12345km" -> "12345"
             "no number here"  -> ""
    """
    raw = (value or "").strip()
    if not raw:
        return ""

    digits: List[str] = []
    started = False
    for ch in raw:
        if ch.isdigit():
            digits.append(ch)
            started = True
            continue
        if started:
            # Commas inside a number are accepted and stripped; any other
            # character ends the number.
            if ch == ",":
                continue
            break
        # Not yet inside a number — keep scanning for the first digit.
    return "".join(digits)


def safe_filename(value: str) -> str:
    value = re.sub(r"[^A-Za-z0-9._ -]+", "", value)
    value = value.strip().replace(" ", "_")
    return value or "export"


def unique_output_path(directory: Path, base_name: str, extension: str) -> Path:
    candidate = directory / f"{base_name}{extension}"
    if not candidate.exists():
        return candidate
    index = 2
    while True:
        candidate = directory / f"{base_name}_{index}{extension}"
        if not candidate.exists():
            return candidate
        index += 1


def set_default_run_font(run, name: str = "Arial", size_pt: int = 12, bold: bool = False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size_pt)
    run.bold = bold


def configure_normal_style(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    style.font.size = Pt(12)


def ordinal_day(n: int) -> str:
    if 10 <= n % 100 <= 20:
        suffix = "th"
    else:
        suffix = {1: "st", 2: "nd", 3: "rd"}.get(n % 10, "th")
    return f"{n}{suffix}"


def format_date_for_rjs(value: str) -> str:
    raw = (value or "").strip()

    def format_dt(dt: datetime) -> str:
        return f"{ordinal_day(dt.day)} {dt.strftime('%B %Y')}"

    if not raw:
        return format_dt(datetime.now())

    patterns = [
        "%d/%m/%Y",
        "%d-%m-%Y",
        "%d/%m/%y",
        "%d-%m-%y",
        "%Y-%m-%d",
        "%d %B %Y",
        "%d %b %Y",
        "%d-%b-%Y",
        "%d-%B-%Y",
        "%d-%b-%y",
        "%d-%B-%y",
    ]
    cleaned = re.sub(r"\b(\d{1,2})(st|nd|rd|th)\b", r"\1", raw, flags=re.IGNORECASE)
    for pattern in patterns:
        try:
            dt = datetime.strptime(cleaned, pattern)
            return format_dt(dt)
        except Exception:
            continue
    return raw


def split_address_lines(value: str) -> List[str]:
    text = (value or "").strip()
    if not text:
        return []

    if "\n" in text:
        parts = [part.strip() for part in text.splitlines() if part.strip()]
    else:
        parts = [part.strip() for part in re.split(r",\s*", text) if part.strip()]
    return parts[:6]


def normalise_rjs_address_block(value: str) -> List[str]:
    """Return exactly 3 address lines for the rigid RJS export layout.

    Rules:
    - Always reserve exactly 3 lines for the address block.
    - If there is 1 line, keep it on line 1 and leave lines 2-3 blank.
    - If there are 2 lines, keep them on lines 1-2 and leave line 3 blank.
    - If there are 3 lines, keep all 3.
    - If there are more than 3 lines, keep the first 2 lines and force the final
      source line onto line 3 so the postcode/last line stays in the expected slot.
    """
    parts = split_address_lines(value)
    if not parts:
        return ["", "", ""]
    if len(parts) == 1:
        return [parts[0], "", ""]
    if len(parts) == 2:
        return [parts[0], parts[1], ""]
    if len(parts) == 3:
        return parts
    return [parts[0], parts[1], parts[-1]]


def build_rjs_docx(output_path: Path, values: Dict[str, str]):
    doc = Document()
    configure_normal_style(doc)

    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.0)

    def add_paragraph(text: str = "", bold: bool = False, center: bool = False, space_after: int = 0) -> None:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_after = Pt(space_after)
        pf.space_before = Pt(0)
        pf.line_spacing = 1.0
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_default_run_font(run, bold=bold)

    # Top address block
    for line in [
        "Collision Engineers Ltd",
        "2 Castle Buildings",
        "147 - 149 Telegraph Road",
        "Heswall",
        "Wirral",
        "United Kingdom",
        "CH60 7SE",
    ]:
        add_paragraph(line)
    add_paragraph(space_after=10)

    add_paragraph("Your Reference:")
    add_paragraph(f"Our Reference: {values.get('reference', '').strip()}")
    add_paragraph("Fee earner: Keeley Garner")
    add_paragraph("Direct dial: 01516650836")
    add_paragraph("Email: k.garner@robertjameslaw.co.uk")
    add_paragraph(format_date_for_rjs(values.get("instruction_date", "")), space_after=22)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("URGENT VEHICLE INSPECTION REQUIRED")
    set_default_run_font(run, size_pt=14, bold=True)

    add_paragraph("Dear Sirs", space_after=14)

    def add_label_value(label: str, value: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r1 = p.add_run(f"{label}:    ")
        set_default_run_font(r1, bold=True)
        r2 = p.add_run(value)
        set_default_run_font(r2, bold=True)

    add_label_value("Our Client", values.get("claimant_name", "").strip())
    add_label_value("Accident", format_date_for_rjs(values.get("incident_date", "").strip()))
    add_label_value("Client vehicle registration", values.get("vrm", "").strip())
    add_label_value("Client vehicle make", "")
    add_label_value("Client vehicle model", values.get("vehicle_model", "").strip())

    add_paragraph(space_after=16)

    claimant = values.get("claimant_name", "the above named client").strip() or "the above named client"
    incident_date = format_date_for_rjs(values.get("incident_date", "").strip())

    narrative_1 = (
        f"I act on behalf of my above named client in the recovery of damages resulting from an accident "
        f"which occurred on {incident_date}."
    )
    narrative_2 = (
        "Please arrange an inspection of my client’s vehicle as soon as possible and provide a report "
        "detailing the damage sustained, costs of repair or cost of replacement if beyond repair."
    )
    narrative_3 = (
        f"I have advised my client of your instruction. Please make arrangements for the inspection with my "
        f"client. {claimant} is available at:"
    )

    add_paragraph(narrative_1, space_after=14)
    add_paragraph(narrative_2, space_after=14)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run(narrative_3)
    set_default_run_font(r, bold=True)

    address_lines = normalise_rjs_address_block(values.get("inspection_address", ""))

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run("Address:    ")
    set_default_run_font(r1)
    r2 = p.add_run(address_lines[0])
    set_default_run_font(r2)

    add_paragraph(address_lines[1])
    add_paragraph(address_lines[2])

    # Keep the RJS layout rigid: always leave exactly 4 empty lines after the
    # 3-line address block before Mobile Tel so the importer reads a stable shape.
    for _ in range(4):
        add_paragraph("")

    add_paragraph("Mobile Tel:")
    add_paragraph(space_after=18)

    boilerplate = (
        "I can confirm that in accordance with the Civil Procedure Rules I have notified the third party of "
        "your involvement in this matter on your behalf and confirmed that I will copy them with your report "
        "once available. Once they are in receipt of your report they may choose to contact you direct with "
        "questions concerning my clients losses. I would be grateful if you could ensure that I receive copies "
        "of any such correspondence along with your replies."
    )
    fees = (
        "Finally having regard to your reasonable fees, I hereby confirm that this firm will be responsible for "
        "the same in accordance with our agreement."
    )

    add_paragraph(boilerplate, space_after=14)
    add_paragraph(fees, space_after=18)
    add_paragraph("Yours faithfully", space_after=40)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run("Keeley Garner")
    set_default_run_font(r1, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r2 = p.add_run("Robert James Solicitors")
    set_default_run_font(r2, bold=True)

    doc.save(str(output_path))


def convert_docx_to_doc_via_word(docx_path: Path, doc_path: Path):
    try:
        import pythoncom  # type: ignore
        from win32com.client import DispatchEx  # type: ignore
    except Exception as exc:
        raise RuntimeError("Microsoft Word automation is not available for DOC export.") from exc

    pythoncom.CoInitialize()
    word = None
    doc = None
    try:
        word = DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(str(docx_path.resolve()), ReadOnly=False, AddToRecentFiles=False, Visible=False)
        # wdFormatDocument = 0 => Word 97-2003 .doc
        doc.SaveAs(str(doc_path.resolve()), FileFormat=0)
    finally:
        if doc is not None:
            try:
                doc.Close(False)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass



def convert_doc_to_docx_via_word(doc_path: Path, docx_path: Path):
    try:
        import pythoncom  # type: ignore
        from win32com.client import DispatchEx  # type: ignore
    except Exception as exc:
        raise RuntimeError("Microsoft Word automation is not available for DOC image extraction.") from exc

    pythoncom.CoInitialize()
    word = None
    doc = None
    try:
        word = DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(str(doc_path.resolve()), ReadOnly=False, AddToRecentFiles=False, Visible=False)
        doc.SaveAs(str(docx_path.resolve()), FileFormat=16)
    finally:
        if doc is not None:
            try:
                doc.Close(False)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass


def convert_doc_to_docx_via_soffice(doc_path: Path, out_dir: Path) -> Path:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        raise RuntimeError("LibreOffice is not installed.")
    command = [
        soffice,
        "--headless",
        "--convert-to",
        "docx",
        "--outdir",
        str(out_dir),
        str(doc_path.resolve()),
    ]
    subprocess.run(command, check=True, capture_output=True)
    converted = out_dir / f"{doc_path.stem}.docx"
    if not converted.exists():
        candidates = list(out_dir.glob("*.docx"))
        if not candidates:
            raise RuntimeError("LibreOffice did not produce a DOCX file.")
        converted = candidates[0]
    return converted


class App:
    def __init__(self, root):
        self.root = root
        self.root.title(APP_TITLE)
        self.current_file: Optional[str] = None
        self.current_source_files: List[str] = []
        self.current_text = ""
        self.current_provider_match: Optional[ProviderMatch] = None
        self.batch_sessions: List[DocumentSession] = []
        self.batch_index = 0
        # Names of fields most recently overwritten by an Engineer Report.
        # Cleared whenever a new document is dragged in. Used to draw a red
        # border on the corresponding value entry in the left panel.
        self.engineer_overridden_fields: set = set()
        # Map of field key -> the value entry/text widget on the left panel.
        # Populated when the values panel is built; used to apply or remove
        # the engineer-override border highlight.
        self.value_entry_widgets: Dict[str, tk.Widget] = {}
        self.last_find_term = ""
        self.find_matches: List[str] = []
        self.find_index = -1
        self.main_container: Optional[ttk.Frame] = None
        self.left_panel: Optional[ttk.Frame] = None
        self.middle_panel: Optional[ttk.Frame] = None
        self.right_panel: Optional[ttk.Frame] = None
        self.toggle_view_button: Optional[ttk.Button] = None
        self.export_doc_button: Optional[tk.Button] = None
        self.export_images_button: Optional[tk.Button] = None
        self.export_json_button: Optional[tk.Button] = None
        self.batch_nav_frame: Optional[ttk.Frame] = None
        self.batch_prev_button: Optional[ttk.Button] = None
        self.batch_next_button: Optional[ttk.Button] = None
        self._settings_save_after_id = None

        self.settings = self.load_settings()
        self.engine = ExtractionEngine(str(DEFAULT_CONFIG_PATH))

        self.values_vars: Dict[str, tk.StringVar] = {key: tk.StringVar() for key in FIELD_KEYS}
        self.mapping_method_vars: Dict[str, tk.StringVar] = {key: tk.StringVar(value=METHOD_CODE_TO_LABEL[DEFAULT_METHOD_BY_FIELD.get(key, "single_label")]) for key in NON_PROVIDER_FIELDS}
        # Mapping config vars cover every field in the rules table,
        # including ``work_provider`` (which has no method dropdown but
        # still needs an editable text value on the right panel).
        self.mapping_config_vars: Dict[str, tk.StringVar] = {key: tk.StringVar() for key in FIELD_KEYS}
        self.mapping_second_config_vars: Dict[str, tk.StringVar] = {key: tk.StringVar() for key in NON_PROVIDER_FIELDS}
        self.mapping_config_frames: Dict[str, ttk.Frame] = {}
        self.mapping_single_entries: Dict[str, ttk.Entry] = {}
        self.mapping_two_label_frames: Dict[str, ttk.Frame] = {}
        self.mapping_fixed_help_labels: Dict[str, tk.Label] = {}
        self.preview_widget: Optional[tk.Text] = None
        self.use_current_date_var = tk.BooleanVar(value=False)
        self.force_postcode_var = tk.BooleanVar(value=False)
        self.engineer_report_var = tk.BooleanVar(value=False)
        self.detected_provider_var = tk.StringVar(value="-")
        self.selected_provider_var = tk.StringVar(value="")
        self.provider_name_var = tk.StringVar(value="")
        self.status_var = tk.StringVar(value="Ready")
        self.find_var = tk.StringVar(value="")
        self.view_mode_var = tk.StringVar(value=self.settings.get("view_mode") or "expanded")
        self.field_highlight_colors = {
            "work_provider": "#ffffff",
            "vrm": "#fff0b3",
            "vehicle_model": "#e9ddff",
            "claimant_name": "#ffd9e6",
            "reference": "#d7f7d0",
            "incident_date": "#ffe2c7",
            "instruction_date": "#d8f0ff",
            "inspection_date": "#f8d4ff",
            "inspection_address": "#e5f5d8",
            "mileage": "#f4e1ff",
            "mileage_unit": "#dff3ff",
            "accident_circumstances": "#ffe3a6",
            "vat_status": "#d6f7ff",
        }

        self.build_ui()
        self.refresh_provider_dropdown()
        self.apply_view_mode(initial=True)
        self.save_settings()

        self.root.bind("<Control-f>", self.focus_find)
        self.root.bind("<F3>", self.find_next)
        self.root.bind("<Shift-F3>", self.find_previous)
        self.root.protocol("WM_DELETE_WINDOW", self.on_close)

    def load_settings(self) -> dict:
        if SETTINGS_PATH.exists():
            try:
                return json.loads(SETTINGS_PATH.read_text(encoding="utf-8"))
            except Exception:
                return {}
        default = {"view_mode": "expanded"}
        SETTINGS_PATH.write_text(json.dumps(default, indent=2, ensure_ascii=False), encoding="utf-8")
        return default

    def save_settings(self):
        SETTINGS_PATH.parent.mkdir(parents=True, exist_ok=True)
        data = {
            "view_mode": self.view_mode_var.get() or "expanded",
        }
        SETTINGS_PATH.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")
        self.settings = data

    def on_close(self):
        self.save_settings()
        self.root.destroy()

    def set_status(self, message: str):
        self.status_var.set(message)

    def apply_notes_to_status(self, notes: List[str]):
        if notes:
            self.status_var.set(" | ".join(notes))

    def is_batch_mode(self) -> bool:
        return len(self.batch_sessions) > 1

    def update_export_ui_for_mode(self):
        is_batch = self.is_batch_mode()
        mappings = [
            (self.export_json_button, "Export JSON", "Batch Export JSON"),
            (self.export_images_button, "Export Images", "Batch Export Images"),
        ]
        for button, single_text, batch_text in mappings:
            if button:
                button.configure(text=batch_text if is_batch else single_text)
        if self.batch_nav_frame:
            if is_batch:
                self.batch_nav_frame.grid()
            else:
                self.batch_nav_frame.grid_remove()

    def save_current_batch_state(self):
        if not self.is_batch_mode():
            return
        if not (0 <= self.batch_index < len(self.batch_sessions)):
            return
        session = self.batch_sessions[self.batch_index]
        session.values = self.collect_values()
        session.selected_provider = self.selected_provider_var.get().strip()
        session.detected_provider_name = self.detected_provider_var.get().strip()
        session.provider_match = self.current_provider_match or session.provider_match

    def show_session(self, index: int):
        if not (0 <= index < len(self.batch_sessions)):
            return
        session = self.batch_sessions[index]
        self.batch_index = index
        self.current_file = session.path
        self.current_source_files = list(session.source_paths or ([session.path] if session.path else []))
        self.current_text = session.text
        self.current_provider_match = session.provider_match

        self.clear_form()
        if self.preview_widget:
            preview_text = self.format_preview_text(session.text[:100000] if session.text else "")
            self.preview_widget.insert("1.0", preview_text)

        self.detected_provider_var.set(session.detected_provider_name or session.provider_match.name)
        for key in FIELD_KEYS:
            self.set_field_value(key, session.values.get(key, ""))

        provider_name = session.selected_provider or session.provider_match.name
        provider_config = self.engine.get_provider_config(provider_name) if provider_name and provider_name != "Unknown / Unmapped" else None
        if provider_config:
            self.load_provider_into_editor(provider_config)
            self.selected_provider_var.set(provider_name)
        else:
            self.clear_provider_editor(keep_provider_dropdown=True)

        self.highlight_extracted_values(session.values)
        base_msg = f"Loaded {Path(session.path).name} ({index + 1}/{len(self.batch_sessions)})"
        if session.notes:
            self.set_status(base_msg + " | " + " | ".join(session.notes))
        else:
            self.set_status(base_msg)
        self.update_export_ui_for_mode()

    def previous_batch_file(self):
        if not self.is_batch_mode():
            return
        self.save_current_batch_state()
        self.show_session((self.batch_index - 1) % len(self.batch_sessions))

    def next_batch_file(self):
        if not self.is_batch_mode():
            return
        self.save_current_batch_state()
        self.show_session((self.batch_index + 1) % len(self.batch_sessions))

    def clear_form(self):
        self.current_source_files = []
        for key in FIELD_KEYS:
            self.set_field_value(key, "")
        self.clear_find_tags()
        self.clear_field_highlights()
        self.clear_engineer_override_highlights()
        if self.preview_widget:
            self.preview_widget.delete("1.0", tk.END)
        self.detected_provider_var.set("-")

    def set_field_value(self, key: str, value: str):
        value = value or ""
        self.values_vars[key].set(value)
        if key == "inspection_address" and hasattr(self, "inspection_address_widget") and self.inspection_address_widget:
            self.inspection_address_widget.delete("1.0", tk.END)
            self.inspection_address_widget.insert("1.0", value)

    def get_field_value(self, key: str) -> str:
        if key == "inspection_address" and hasattr(self, "inspection_address_widget") and self.inspection_address_widget:
            return self.inspection_address_widget.get("1.0", "end-1c").rstrip()
        return self.values_vars[key].get().strip()

    def _sync_inspection_address_var_from_widget(self, _event=None):
        if hasattr(self, "inspection_address_widget") and self.inspection_address_widget:
            self.values_vars["inspection_address"].set(self.inspection_address_widget.get("1.0", "end-1c").rstrip())

    def build_ui(self):
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(1, weight=1)

        top_bar = ttk.Frame(self.root)
        top_bar.grid(row=0, column=0, sticky="ew", padx=10, pady=(10, 0))
        top_bar.columnconfigure(0, weight=1)
        self.toggle_view_button = ttk.Button(top_bar, text="⤡", width=3, command=self.toggle_view_mode)
        self.toggle_view_button.grid(row=0, column=1, sticky="e")

        self.main_container = ttk.Frame(self.root)
        self.main_container.grid(row=1, column=0, sticky="nsew", padx=10, pady=10)
        self.main_container.rowconfigure(0, weight=1)
        for col in range(3):
            self.main_container.columnconfigure(col, weight=1, uniform="panels")

        left = ttk.Frame(self.main_container, padding=8)
        middle = ttk.Frame(self.main_container, padding=8)
        right = ttk.Frame(self.main_container, padding=8)
        self.left_panel = left
        self.middle_panel = middle
        self.right_panel = right

        left.grid(row=0, column=0, sticky="nsew")
        middle.grid(row=0, column=1, sticky="nsew")
        right.grid(row=0, column=2, sticky="nsew")

        self.build_values_panel(left)
        self.build_preview_panel(middle)
        self.build_mapping_panel(right)

        status = ttk.Label(self.root, textvariable=self.status_var, relief="sunken", anchor="w")
        status.grid(row=2, column=0, sticky="ew")

        self.register_drop_targets()

    def update_view_toggle_button(self):
        if not self.toggle_view_button:
            return
        if self.view_mode_var.get() == "minimal":
            self.toggle_view_button.configure(text="⤢")
        else:
            self.toggle_view_button.configure(text="⤡")

    def _set_fixed_window_geometry(self, geometry: str):
        self.root.geometry(geometry)
        try:
            size_part = geometry.split("+", 1)[0]
            width, height = size_part.split("x", 1)
            width_i, height_i = int(width), int(height)
            self.root.minsize(width_i, height_i)
            self.root.maxsize(width_i, height_i)
            self.root.resizable(False, False)
        except Exception:
            pass

    def apply_view_mode(self, initial: bool = False):
        if not self.main_container or not self.left_panel or not self.middle_panel or not self.right_panel:
            return

        mode = self.view_mode_var.get() or "expanded"

        # Always reset panel placement explicitly before applying the target layout.
        # This avoids stale grid metadata such as an old columnspan carrying over
        # after switching between views.
        self.left_panel.grid_forget()
        self.middle_panel.grid_forget()
        self.right_panel.grid_forget()

        if mode == "minimal":
            self.left_panel.grid(row=0, column=0, columnspan=1, sticky="nsew")
            self.main_container.columnconfigure(0, weight=1, uniform="panels", minsize=380)
            self.main_container.columnconfigure(1, weight=0, uniform="panels", minsize=0)
            self.main_container.columnconfigure(2, weight=0, uniform="panels", minsize=0)
            self._set_fixed_window_geometry(MINIMAL_WINDOW_GEOMETRY)
        else:
            self.left_panel.grid(row=0, column=0, columnspan=1, sticky="nsew")
            self.middle_panel.grid(row=0, column=1, columnspan=1, sticky="nsew")
            self.right_panel.grid(row=0, column=2, columnspan=1, sticky="nsew")
            for col in range(3):
                self.main_container.columnconfigure(col, weight=1, uniform="panels", minsize=440)
            self._set_fixed_window_geometry(EXPANDED_WINDOW_GEOMETRY)

        self.main_container.update_idletasks()
        self.update_view_toggle_button()
        self.update_export_ui_for_mode()
        if not initial:
            self.save_settings()

    def toggle_view_mode(self):
        self.view_mode_var.set("minimal" if self.view_mode_var.get() == "expanded" else "expanded")
        self.apply_view_mode()

    def on_export_doc_click(self):
        try:
            if self.is_batch_mode():
                self.batch_export(export_doc=True, export_images=False, export_json=False)
            else:
                self.export_doc(export_doc=True, export_images=False)
        except Exception as exc:
            messagebox.showerror(APP_TITLE, f"Export failed: {exc}")
            self.set_status(f"Export failed: {exc}")

    def on_export_images_click(self):
        try:
            if self.is_batch_mode():
                self.batch_export(export_doc=False, export_images=True, export_json=False)
            else:
                self.export_doc(export_doc=False, export_images=True)
        except Exception as exc:
            messagebox.showerror(APP_TITLE, f"Export failed: {exc}")
            self.set_status(f"Export failed: {exc}")

    def on_export_json_click(self):
        try:
            if self.is_batch_mode():
                self.batch_export(export_doc=False, export_images=False, export_json=True)
            else:
                self.export_json_string()
        except Exception as exc:
            messagebox.showerror(APP_TITLE, f"Export failed: {exc}")
            self.set_status(f"Export failed: {exc}")
    def register_drop_targets(self):
        if not HAS_DND:
            return

        widgets = [self.root, self.main_container, self.left_panel, self.middle_panel, self.right_panel]
        seen = set()
        for widget in widgets:
            if not widget:
                continue
            widget_id = str(widget)
            if widget_id in seen:
                continue
            seen.add(widget_id)
            try:
                widget.drop_target_register(DND_FILES)
                widget.dnd_bind("<<Drop>>", self.on_drop)
            except Exception:
                pass

    def build_values_panel(self, parent):
        parent.columnconfigure(0, weight=0, minsize=120)
        parent.columnconfigure(1, weight=1, minsize=180)
        row = 0

        header = ttk.Frame(parent)
        header.grid(row=row, column=0, columnspan=2, sticky="ew", pady=(0, 8))
        header.columnconfigure(1, weight=1)
        ttk.Label(header, text="Detected Provider:").grid(row=0, column=0, sticky="w")
        ttk.Label(header, textvariable=self.detected_provider_var).grid(row=0, column=1, sticky="w", padx=(6, 0))
        row += 1

        self.inspection_address_widget = None
        for key, label in DEFAULT_FIELDS:
            color = "#ffffff" if key == "work_provider" else self.field_highlight_colors.get(key, "#ffffff")
            tk.Label(parent, text=label, bg=color, anchor="w", padx=4).grid(row=row, column=0, sticky="ew", pady=4, padx=(0, 4))
            if key == "inspection_address":
                entry = tk.Text(parent, bg="white", relief="solid", bd=1, width=1, height=6, wrap="word",
                                highlightthickness=0, highlightbackground="white", highlightcolor="white")
                entry.grid(row=row, column=1, sticky="ew", pady=4)
                entry.bind("<KeyRelease>", self._sync_inspection_address_var_from_widget)
                entry.bind("<FocusOut>", self._sync_inspection_address_var_from_widget)
                self.inspection_address_widget = entry
            else:
                entry = tk.Entry(parent, textvariable=self.values_vars[key], bg="white", relief="solid", bd=1, width=1,
                                 highlightthickness=0, highlightbackground="white", highlightcolor="white")
                entry.grid(row=row, column=1, sticky="ew", pady=4)
            self.value_entry_widgets[key] = entry
            row += 1

        buttons = ttk.Frame(parent)
        buttons.grid(row=row, column=0, columnspan=2, sticky="ew", pady=(12, 0))
        buttons.columnconfigure(0, weight=1)
        self.export_json_button = tk.Button(buttons, text="Export JSON", command=self.on_export_json_click)
        self.export_json_button.grid(row=0, column=0, sticky="ew", pady=(0, 6))
        self.export_images_button = tk.Button(buttons, text="Export Images", command=self.on_export_images_click)
        self.export_images_button.grid(row=1, column=0, sticky="ew", pady=(0, 0))
        self.export_doc_button = None

        self.batch_nav_frame = ttk.Frame(parent)
        self.batch_nav_frame.grid(row=row + 1, column=0, columnspan=2, sticky="ew", pady=(12, 0))
        self.batch_nav_frame.columnconfigure(0, weight=1)
        nav_inner = ttk.Frame(self.batch_nav_frame)
        nav_inner.grid(row=0, column=0)
        self.batch_prev_button = ttk.Button(nav_inner, text="←", width=3, command=self.previous_batch_file)
        self.batch_prev_button.pack(side="left", padx=(0, 6))
        self.batch_next_button = ttk.Button(nav_inner, text="→", width=3, command=self.next_batch_file)
        self.batch_next_button.pack(side="left")
        self.batch_nav_frame.grid_remove()

    def build_preview_panel(self, parent):
        parent.columnconfigure(0, weight=1)
        parent.rowconfigure(1, weight=1)

        search_bar = ttk.Frame(parent)
        search_bar.grid(row=0, column=0, sticky="ew")
        search_bar.columnconfigure(1, weight=1)

        ttk.Label(search_bar, text="Source Text Preview").grid(row=0, column=0, sticky="w")
        self.find_entry = ttk.Entry(search_bar, textvariable=self.find_var)
        self.find_entry.grid(row=0, column=1, sticky="ew", padx=(8, 4))
        self.find_entry.bind("<Return>", self.find_next)
        ttk.Button(search_bar, text="Find Next", command=self.find_next).grid(row=0, column=2, padx=2)
        ttk.Button(search_bar, text="Find Previous", command=self.find_previous).grid(row=0, column=3, padx=2)

        preview_frame = ttk.Frame(parent)
        preview_frame.grid(row=1, column=0, sticky="nsew", pady=(4, 0))
        preview_frame.columnconfigure(0, weight=1)
        preview_frame.rowconfigure(0, weight=1)

        preview = tk.Text(preview_frame, wrap="none")
        preview.grid(row=0, column=0, sticky="nsew")
        preview.tag_configure("find_all", background="#fff2a8")
        preview.tag_configure("find_current", background="#ffb347")
        for field_name, color in self.field_highlight_colors.items():
            preview.tag_configure(f"field_{field_name}", background=color)
        preview_scroll = ttk.Scrollbar(preview_frame, orient="vertical", command=preview.yview)
        preview_scroll.grid(row=0, column=1, sticky="ns")
        preview_xscroll = ttk.Scrollbar(preview_frame, orient="horizontal", command=preview.xview)
        preview_xscroll.grid(row=1, column=0, sticky="ew")
        preview.configure(yscrollcommand=preview_scroll.set, xscrollcommand=preview_xscroll.set)
        preview.bind("<Control-c>", self.copy_preview_selection)
        preview.bind("<Control-C>", self.copy_preview_selection)
        preview.bind("<<Copy>>", self.copy_preview_selection)
        self.preview_widget = preview

    def build_mapping_panel(self, parent):
        parent.columnconfigure(0, weight=1)
        row = 0

        ttk.Label(parent, text="Provider Setup").grid(row=row, column=0, sticky="w")
        row += 1

        form = ttk.Frame(parent)
        form.grid(row=row, column=0, sticky="nsew")
        form.columnconfigure(1, weight=1)
        row += 1

        ttk.Label(form, text="Saved Provider").grid(row=0, column=0, sticky="w", pady=4)
        self.provider_combo = ttk.Combobox(form, textvariable=self.selected_provider_var, state="readonly")
        self.provider_combo.grid(row=0, column=1, sticky="ew", pady=4)
        self.provider_combo.bind("<<ComboboxSelected>>", self.on_provider_selected)

        ttk.Label(form, text="Provider Name").grid(row=1, column=0, sticky="w", pady=4)
        ttk.Entry(form, textvariable=self.provider_name_var).grid(row=1, column=1, sticky="ew", pady=4)

        ttk.Label(form, text="Detect this provider when these phrases appear\n(one per line)").grid(
            row=2, column=0, sticky="nw", pady=4
        )
        self.detect_box = tk.Text(form, height=5, wrap="word")
        self.detect_box.grid(row=2, column=1, sticky="ew", pady=4)

        self.engineer_report_check = tk.Checkbutton(
            form,
            text="Engineer Report",
            variable=self.engineer_report_var,
            bg=self.root.cget("bg"),
            anchor="w",
        )
        self.engineer_report_check.grid(row=3, column=1, sticky="w", pady=4)

        ttk.Separator(parent, orient="horizontal").grid(row=row, column=0, sticky="ew", pady=8)
        row += 1

        rules_frame = ttk.Frame(parent)
        rules_frame.grid(row=row, column=0, sticky="nsew")
        rules_frame.columnconfigure(0, minsize=120)
        rules_frame.columnconfigure(1, minsize=120)
        rules_frame.columnconfigure(2, weight=1)
        row += 1

        ttk.Label(rules_frame, text="Field").grid(row=0, column=0, sticky="w")
        ttk.Label(rules_frame, text="Method").grid(row=0, column=1, sticky="w")
        ttk.Label(rules_frame, text="Config").grid(row=0, column=2, sticky="w")

        # The Work Provider row sits at the top of the rules table. It
        # has no method dropdown — the user just types a literal value
        # that ends up in the Detected Fields panel and JSON export. The
        # rest of the rows mirror the order of DEFAULT_FIELDS so the
        # right panel reads in the same sequence as the left panel.
        work_provider_row = 1
        tk.Label(rules_frame, text="Work Provider", bg="#ffffff", anchor="w", padx=4).grid(
            row=work_provider_row, column=0, sticky="ew", pady=3, padx=(0, 6)
        )
        # No widget in the Method column for this row — the entry takes
        # the place of the config widget below. (Leaving the column
        # empty keeps the header alignment intact.)
        wp_config_frame = tk.Frame(rules_frame, bg="white")
        wp_config_frame.grid(row=work_provider_row, column=2, sticky="ew", pady=3)
        wp_config_frame.columnconfigure(0, weight=1)
        self.mapping_config_frames["work_provider"] = wp_config_frame
        wp_entry = tk.Entry(
            wp_config_frame,
            textvariable=self.mapping_config_vars["work_provider"],
            bg="white",
            relief="solid",
            bd=1,
        )
        wp_entry.grid(row=0, column=0, sticky="ew")
        self.mapping_single_entries["work_provider"] = wp_entry

        ordered_fields = [
            ("vrm", "VRM"),
            ("vehicle_model", "Vehicle Model"),
            ("claimant_name", "Claimant Name"),
            ("reference", "Reference"),
            ("incident_date", "Incident Date"),
            ("instruction_date", "Instruction Date"),
            ("inspection_date", "Inspection Date"),
            ("inspection_address", "Inspection Address"),
            ("accident_circumstances", "Accident Circumstances"),
            ("vat_status", "VAT Status"),
            ("mileage", "Mileage"),
            ("mileage_unit", "Mileage Unit"),
        ]

        method_labels = [label for _, label in METHOD_CHOICES]
        table_row = 2  # row 0 = column headers, row 1 = Work Provider
        for key, label in ordered_fields:
            color = self.field_highlight_colors.get(key, "#ffffff")
            tk.Label(rules_frame, text=label, bg=color, anchor="w", padx=4).grid(row=table_row, column=0, sticky="ew", pady=3, padx=(0, 6))
            combo = ttk.Combobox(
                rules_frame,
                textvariable=self.mapping_method_vars[key],
                values=method_labels,
                state="readonly",
                width=14,
            )
            combo.grid(row=table_row, column=1, sticky="w", pady=3, padx=(0, 6))

            # Presence-check fields (VAT Status, Mileage Unit) don't use a
            # mapping method — the user just enters comma-separated tokens
            # describing the positive scenario. Hide the dropdown and force
            # the underlying method to single_label so save/load is stable.
            if key in PRESENCE_CHECK_FIELDS:
                combo.grid_remove()
                self.mapping_method_vars[key].set(METHOD_CODE_TO_LABEL["single_label"])

            config_frame = tk.Frame(rules_frame, bg="white")
            config_frame.grid(row=table_row, column=2, sticky="ew", pady=3)
            config_frame.columnconfigure(0, weight=1)
            self.mapping_config_frames[key] = config_frame

            single_entry = tk.Entry(config_frame, textvariable=self.mapping_config_vars[key], bg="white", relief="solid", bd=1)
            single_entry.grid(row=0, column=0, sticky="ew")
            self.mapping_single_entries[key] = single_entry

            two_frame = tk.Frame(config_frame, bg="white")
            two_frame.columnconfigure(0, weight=1)
            two_frame.columnconfigure(1, weight=1)
            tk.Entry(two_frame, textvariable=self.mapping_config_vars[key], bg="white", relief="solid", bd=1).grid(row=0, column=0, sticky="ew", padx=(0, 4))
            tk.Entry(two_frame, textvariable=self.mapping_second_config_vars[key], bg="white", relief="solid", bd=1).grid(row=0, column=1, sticky="ew")
            self.mapping_two_label_frames[key] = two_frame

            self.mapping_fixed_help_labels[key] = None

            if key == "inspection_date":
                current_date_chk = tk.Checkbutton(
                    config_frame,
                    text="Use Current Date",
                    variable=self.use_current_date_var,
                    bg="white",
                    anchor="w",
                    command=self.on_use_current_date_toggled,
                )
                current_date_chk.grid(row=0, column=1, sticky="w", padx=(8, 0))
                self.inspection_date_checkbox = current_date_chk

            if key == "inspection_address":
                force_postcode_chk = tk.Checkbutton(
                    config_frame,
                    text="Force Postcode",
                    variable=self.force_postcode_var,
                    bg="white",
                    anchor="w",
                )
                force_postcode_chk.grid(row=0, column=1, sticky="w", padx=(8, 0))
                self.force_postcode_checkbox = force_postcode_chk

            combo.bind("<<ComboboxSelected>>", lambda _event, field_key=key: self.update_method_ui(field_key))
            self.mapping_method_vars[key].trace_add("write", lambda *_args, field_key=key: self.update_method_ui(field_key))
            self.update_method_ui(key)
            table_row += 1

        buttons = ttk.Frame(parent)
        buttons.grid(row=row, column=0, sticky="ew", pady=(12, 0))
        ttk.Button(buttons, text="Use this Provider", command=self.use_selected_provider).pack(side="left")
        right_buttons = ttk.Frame(buttons)
        right_buttons.pack(side="right")
        ttk.Button(right_buttons, text="Delete Mapping", command=self.delete_provider_mapping).pack(side="left", padx=(0, 6))
        ttk.Button(right_buttons, text="Save Mapping", command=self.save_provider_mapping).pack(side="left")


    def on_use_current_date_toggled(self):
        self.update_method_ui("inspection_date")

    def update_method_ui(self, field_key: str):
        frame = self.mapping_config_frames.get(field_key)
        if not frame:
            return
        method_label = self.mapping_method_vars[field_key].get().strip()
        method_code = METHOD_LABEL_TO_CODE.get(method_label, DEFAULT_METHOD_BY_FIELD.get(field_key, "single_label"))

        # Presence-check fields (VAT Status, Mileage Unit) ignore the mapping
        # method entirely — the user just enters comma-separated tokens.
        if field_key in PRESENCE_CHECK_FIELDS:
            method_code = "single_label"

        single_entry = self.mapping_single_entries.get(field_key)
        two_frame = self.mapping_two_label_frames.get(field_key)
        help_label = self.mapping_fixed_help_labels.get(field_key)

        if single_entry:
            single_entry.grid_remove()
        if two_frame:
            two_frame.grid_remove()
        if help_label:
            help_label.grid_remove()

        if method_code in ("two_labels", "fixed_position_label", "single_label_offset"):
            if two_frame:
                two_frame.grid(row=0, column=0, sticky="ew")
        else:
            if single_entry:
                single_entry.grid(row=0, column=0, sticky="ew")

        if field_key == "inspection_date":
            disable_inputs = bool(self.use_current_date_var.get())
            state = "disabled" if disable_inputs else "normal"
            if single_entry:
                try:
                    single_entry.configure(state=state)
                except Exception:
                    pass
            if two_frame:
                for child in two_frame.winfo_children():
                    try:
                        child.configure(state=state)
                    except Exception:
                        pass

    def format_preview_text(self, text: str) -> str:
        lines = text.splitlines()
        if not lines:
            return ""
        width = max(4, len(str(len(lines))))
        return "\n".join(f"{idx:0{width}d} | {line}" for idx, line in enumerate(lines, start=1))
    def refresh_provider_dropdown(self):
        names = self.engine.list_provider_names()
        self.provider_combo["values"] = names

    def parse_drop_files(self, raw: str) -> List[str]:
        raw = (raw or "").strip()
        if not raw:
            return []
        try:
            items = list(self.root.tk.splitlist(raw))
        except Exception:
            items = [raw]
        paths: List[str] = []
        for item in items:
            item = item.strip()
            if item.startswith("{") and item.endswith("}"):
                item = item[1:-1]
            if item:
                paths.append(item)
        return paths

    def on_drop(self, event):
        paths = self.parse_drop_files(getattr(event, "data", ""))
        if paths:
            self.load_files(paths)

    def rescan_with_provider(self, provider: ProviderMatch):
        values, extract_notes = self.engine.extract_fields(self.current_text, provider)
        self.current_provider_match = provider
        self.detected_provider_var.set(provider.name)

        for key in FIELD_KEYS:
            self.set_field_value(key, values.get(key, ""))

        self.load_provider_into_editor(provider.config)
        self.highlight_extracted_values(values)
        if extract_notes:
            self.apply_notes_to_status(extract_notes)
        else:
            self.set_status(f"Using provider: {provider.name}")

        if self.is_batch_mode() and 0 <= self.batch_index < len(self.batch_sessions):
            session = self.batch_sessions[self.batch_index]
            session.values = dict(values)
            session.provider_match = provider
            session.selected_provider = provider.name
            session.detected_provider_name = provider.name
            session.notes = list(extract_notes)

    def is_engineer_report_provider(self, provider: ProviderMatch) -> bool:
        return bool((provider.config or {}).get("engineer_report"))

    def apply_engineer_report_session(self, session: DocumentSession):
        if not self.current_file or not self.collect_values().get("work_provider", "").strip():
            messagebox.showerror(APP_TITLE, "You must process an instruction before an engineer's report")
            self.set_status("You must process an instruction before an engineer's report")
            return

        base_values = self.collect_values()
        overrides = self._compute_engineer_overrides(base_values, session.values)

        merged_values = dict(base_values)
        for key in NON_PROVIDER_FIELDS:
            new_value = (session.values.get(key) or "").strip()
            if new_value:
                merged_values[key] = new_value

        if self.preview_widget:
            self.preview_widget.delete("1.0", tk.END)
            preview_text = self.format_preview_text(session.text[:100000] if session.text else "")
            self.preview_widget.insert("1.0", preview_text)

        for key in FIELD_KEYS:
            self.set_field_value(key, merged_values.get(key, ""))

        self.current_file = session.path
        merged_sources = list(self.current_source_files or [])
        for src in (session.source_paths or [session.path]):
            if src not in merged_sources:
                merged_sources.append(src)
        self.current_source_files = merged_sources
        self.current_text = session.text
        self.current_provider_match = session.provider_match
        self.detected_provider_var.set(session.detected_provider_name or session.provider_match.name)

        if session.provider_match.name != "Unknown / Unmapped":
            self.load_provider_into_editor(session.provider_match.config)
            self.selected_provider_var.set(session.provider_match.name)
        else:
            self.clear_provider_editor(keep_provider_dropdown=True)

        self.highlight_extracted_values(merged_values)
        self.engineer_overridden_fields = overrides
        self.apply_engineer_override_highlights()
        status = f"Applied engineer report: {Path(session.path).name}"
        if session.notes:
            status += " | " + " | ".join(session.notes)
        self.set_status(status)

    def combine_instruction_and_engineer_report(self, instruction_session: DocumentSession, engineer_session: DocumentSession) -> DocumentSession:
        merged_values = dict(instruction_session.values)
        for key in NON_PROVIDER_FIELDS:
            new_value = (engineer_session.values.get(key) or "").strip()
            if new_value:
                merged_values[key] = new_value
        merged_notes = list(instruction_session.notes) + [f"Applied engineer report: {Path(engineer_session.path).name}"] + list(engineer_session.notes)
        return DocumentSession(
            path=engineer_session.path,
            text=engineer_session.text,
            provider_match=engineer_session.provider_match,
            values=merged_values,
            notes=merged_notes,
            selected_provider=engineer_session.provider_match.name if engineer_session.provider_match.name != "Unknown / Unmapped" else instruction_session.selected_provider,
            detected_provider_name=engineer_session.provider_match.name or engineer_session.detected_provider_name,
            source_paths=list((instruction_session.source_paths or [instruction_session.path])) + [src for src in (engineer_session.source_paths or [engineer_session.path]) if src not in (instruction_session.source_paths or [instruction_session.path])],
        )

    def process_file_to_session(self, path: str) -> DocumentSession:
        text, read_notes = self.engine.extract_text(path)
        provider = self.engine.detect_provider(text)
        values, extract_notes = self.engine.extract_fields(text, provider)
        combined_notes = read_notes + extract_notes
        selected_provider = provider.name if provider.name != "Unknown / Unmapped" else ""
        detected_provider_name = provider.name
        return DocumentSession(
            path=path,
            text=text,
            provider_match=provider,
            values=values,
            notes=combined_notes,
            selected_provider=selected_provider,
            detected_provider_name=detected_provider_name,
            source_paths=[path],
        )

    def load_file(self, path: str):
        self.load_files([path])

    def load_files(self, paths: List[str]):
        clean_paths = [str(Path(path)) for path in paths if str(path).strip()]
        if not clean_paths:
            return

        self.save_current_batch_state()
        self.set_status("Reading document...")

        sessions: List[DocumentSession] = []
        failures: List[str] = []
        for path in clean_paths:
            try:
                sessions.append(self.process_file_to_session(path))
            except Exception as exc:
                failures.append(f"{Path(path).name}: {exc}")

        if not sessions:
            messagebox.showerror(APP_TITLE, "No documents could be loaded.\n\n" + "\n".join(failures[:10]))
            self.set_status("No documents could be loaded.")
            self.batch_sessions = []
            self.batch_index = 0
            self.current_file = None
            self.current_source_files = []
            self.current_text = ""
            self.current_provider_match = None
            self.update_export_ui_for_mode()
            return

        engineer_sessions = [s for s in sessions if self.is_engineer_report_provider(s.provider_match)]
        instruction_sessions = [s for s in sessions if not self.is_engineer_report_provider(s.provider_match)]

        if len(clean_paths) == 1 and len(engineer_sessions) == 1 and not instruction_sessions:
            if self.is_batch_mode():
                messagebox.showerror(APP_TITLE, "You cannot process multiple documents with an engineer's report")
                self.set_status("You cannot process multiple documents with an engineer's report")
                return
            if not self.current_file or not self.get_field_value("work_provider").strip():
                messagebox.showerror(APP_TITLE, "You must process an instruction before an engineer's report")
                self.set_status("You must process an instruction before an engineer's report")
                return
            self.apply_engineer_report_session(engineer_sessions[0])
            return

        if engineer_sessions:
            if len(engineer_sessions) > 1:
                messagebox.showerror(APP_TITLE, "You cannot process more than one engineer's report at once")
                self.set_status("You cannot process more than one engineer's report at once")
                return
            if not instruction_sessions:
                messagebox.showerror(APP_TITLE, "You must process an instruction before an engineer's report")
                self.set_status("You must process an instruction before an engineer's report")
                return
            if len(instruction_sessions) > 1:
                messagebox.showerror(APP_TITLE, "You cannot process multiple documents with an engineer's report")
                self.set_status("You cannot process multiple documents with an engineer's report")
                return

            self.batch_sessions = []
            self.batch_index = 0
            self.clear_form()
            instruction_session = instruction_sessions[0]
            engineer_session = engineer_sessions[0]
            overrides = self._compute_engineer_overrides(
                instruction_session.values, engineer_session.values
            )
            merged = self.combine_instruction_and_engineer_report(instruction_session, engineer_session)
            self.current_file = merged.path
            self.current_source_files = list(merged.source_paths or [])
            self.current_text = merged.text
            self.current_provider_match = merged.provider_match

            if self.preview_widget:
                preview_text = self.format_preview_text(merged.text[:100000] if merged.text else "")
                self.preview_widget.insert("1.0", preview_text)

            self.detected_provider_var.set(merged.detected_provider_name)
            for key in FIELD_KEYS:
                self.set_field_value(key, merged.values.get(key, ""))

            if merged.provider_match.name != "Unknown / Unmapped":
                self.load_provider_into_editor(merged.provider_match.config)
                self.selected_provider_var.set(merged.provider_match.name)
            else:
                self.clear_provider_editor(keep_provider_dropdown=True)

            self.highlight_extracted_values(merged.values)
            self.engineer_overridden_fields = overrides
            self.apply_engineer_override_highlights()
            if merged.notes:
                self.apply_notes_to_status(merged.notes)
            else:
                self.set_status("Document loaded")
            self.update_export_ui_for_mode()
            if failures:
                messagebox.showwarning(APP_TITLE, "Some documents could not be loaded:\n\n" + "\n".join(failures[:10]))
            return

        if len(sessions) == 1:
            session = sessions[0]
            self.batch_sessions = []
            self.batch_index = 0
            self.clear_form()
            self.current_file = session.path
            self.current_source_files = list(session.source_paths or [session.path])
            self.current_text = session.text
            self.current_provider_match = session.provider_match

            if self.preview_widget:
                preview_text = self.format_preview_text(session.text[:100000] if session.text else "")
                self.preview_widget.insert("1.0", preview_text)

            self.detected_provider_var.set(session.detected_provider_name)
            for key in FIELD_KEYS:
                self.set_field_value(key, session.values.get(key, ""))

            if session.provider_match.name != "Unknown / Unmapped":
                self.load_provider_into_editor(session.provider_match.config)
                self.selected_provider_var.set(session.provider_match.name)
            else:
                self.clear_provider_editor(keep_provider_dropdown=True)

            self.highlight_extracted_values(session.values)
            if session.notes:
                self.apply_notes_to_status(session.notes)
            else:
                self.set_status("Document loaded")
            self.update_export_ui_for_mode()
            if failures:
                messagebox.showwarning(APP_TITLE, "Some documents could not be loaded:\n\n" + "\n".join(failures[:10]))
            return

        self.clear_form()
        self.batch_sessions = sessions
        self.batch_index = 0
        self.show_session(0)

        if failures:
            messagebox.showwarning(APP_TITLE, "Some documents could not be loaded:\n\n" + "\n".join(failures[:10]))

    def collect_values(self) -> Dict[str, str]:
        return {key: self.get_field_value(key).strip() for key in FIELD_KEYS}

    def prepare_export_values(self, values: Dict[str, str], provider_name: Optional[str] = None) -> Dict[str, str]:
        prepared = dict(values)
        chosen_provider_name = (provider_name or self.selected_provider_var.get() or self.detected_provider_var.get() or "").strip()
        provider_config = self.engine.get_provider_config(chosen_provider_name) if chosen_provider_name and chosen_provider_name != "Unknown / Unmapped" else None
        force_postcode = bool((provider_config or {}).get("force_postcode_for_inspection_address"))
        prepared["vrm"] = post_process_extracted_value("vrm", prepared.get("vrm", ""))
        prepared["inspection_address"] = post_process_extracted_value("inspection_address", prepared.get("inspection_address", ""), force_postcode=force_postcode)
        prepared["mileage"] = post_process_extracted_value("mileage", prepared.get("mileage", ""))
        # Normalise dates to DD/MM/YYYY at export time. The Detected
        # Fields panel preserves whatever shape the document used; the
        # JSON output is always canonicalised so downstream management
        # software's import schema is satisfied. Unparseable values are
        # left as-is so the user notices the issue at import time.
        for date_field in ("incident_date", "instruction_date", "inspection_date"):
            prepared[date_field] = normalise_date_value(prepared.get(date_field, ""))
        return prepared

    def export_base_name(self, values: Dict[str, str], source_path: Optional[str] = None) -> str:
        work_provider_value = values.get("work_provider", "").strip()
        vrm_raw = (values.get("vrm") or "").strip()

        # When there is no Work Provider, fall back to the source file's
        # stem so the output filenames stay tied to which document the
        # images / data came from (e.g. ``Images_img_1.jpeg`` rather than
        # ``UnknownVRM_img_1.jpeg``). This is reachable in practice only
        # for Image export, since JSON export refuses to run when Work
        # Provider is empty.
        if not work_provider_value:
            if source_path:
                source_stem = Path(source_path).stem.strip()
                if source_stem:
                    return safe_filename(source_stem)
            return safe_filename(vrm_raw or "UnknownVRM")

        vrm = safe_filename(vrm_raw or "UnknownVRM")
        provider_slug = safe_filename(work_provider_value)
        return f"{provider_slug}_{vrm}"

    def _display_name_for_alert(self, source_file: str) -> str:
        return Path(source_file).name

    def export_outputs_for_values(self, source_file, values: Dict[str, str], export_doc: bool = True, export_images: bool = False, export_json: bool = False, provider_name: Optional[str] = None) -> Dict[str, object]:
        values = self.prepare_export_values(values, provider_name=provider_name)

        # Pick the primary source path (first if a list) so the base name
        # can fall back to the source filename when Work Provider is empty.
        if isinstance(source_file, (list, tuple)):
            primary_source = next((str(s) for s in source_file if s), None)
        else:
            primary_source = str(source_file) if source_file else None

        base_name = self.export_base_name(values, source_path=primary_source)
        result: Dict[str, object] = {
            "doc_path": None,
            "json_path": None,
            "image_count": 0,
            "base_name": base_name,
        }

        if export_doc:
            output_doc = unique_output_path(OUTPUT_DIR, base_name, ".doc")
            fallback_docx = output_doc.with_suffix(".docx")
            with tempfile.TemporaryDirectory() as tmpdir:
                temp_docx = Path(tmpdir) / f"{base_name}.docx"
                build_rjs_docx(temp_docx, values)
                try:
                    convert_docx_to_doc_via_word(temp_docx, output_doc)
                    result["doc_path"] = output_doc
                except Exception:
                    shutil.copy2(temp_docx, fallback_docx)
                    result["doc_path"] = fallback_docx

        if export_images:
            source_files = source_file if isinstance(source_file, (list, tuple)) else [source_file]
            total_saved = 0
            for src in source_files:
                saved, _image_notes = self.engine.extract_images_to_desktop(str(src), OUTPUT_DIR, base_name)
                total_saved += len(saved)
            result["image_count"] = total_saved

        if export_json:
            ordered = {FIELD_LABELS[key]: values.get(key, "") for key in FIELD_KEYS}
            json_string = json.dumps(ordered, ensure_ascii=False, indent=2)
            output_json = unique_output_path(OUTPUT_DIR, base_name, ".json")
            output_json.write_text(json_string, encoding="utf-8")
            result["json_path"] = output_json
            result["json_string"] = json_string

        return result

    def export_doc(self, export_doc: bool = True, export_images: bool = False):
        if not self.current_file:
            messagebox.showerror(APP_TITLE, "Please drag in a PDF, DOCX, DOC, EML, or MSG first.")
            return

        prepared_values = self.prepare_export_values(self.collect_values())

        # Image export is allowed even when Work Provider is empty —
        # the filename falls back to the source document's stem so the
        # output stays meaningful (e.g. ``Images_img_1.jpeg``).
        # Other export types remain gated to avoid stray
        # ``UnknownVRM_*`` outputs for unidentified documents.
        if not export_images and not (prepared_values.get("work_provider") or "").strip():
            return

        self.export_outputs_for_values(
            source_file=self.current_source_files or [self.current_file],
            values=prepared_values,
            export_doc=export_doc,
            export_images=export_images,
            export_json=False,
            provider_name=(self.selected_provider_var.get() or self.detected_provider_var.get() or "").strip(),
        )

        if export_doc:
            self.set_status("Exported DOC to Desktop")
        elif export_images:
            self.set_status("Exported Images to Desktop")
        else:
            self.set_status("Nothing exported")

    def batch_export(self, export_doc: bool = False, export_images: bool = False, export_json: bool = False):
        if not self.is_batch_mode():
            return

        self.save_current_batch_state()
        files_processed = 0
        total_images = 0
        failures: List[str] = []

        for session in self.batch_sessions:
            # Skip unidentified sessions for JSON / DOC exports so we
            # never produce ``UnknownVRM.*`` outputs the user can't
            # match to their docs. Image export is exempt — it falls
            # back to the source filename as a base name, so each
            # session's images are still meaningfully named.
            if not export_images and not (session.values.get("work_provider") or "").strip():
                continue
            try:
                result = self.export_outputs_for_values(
                    source_file=session.source_paths or [session.path],
                    values=session.values,
                    export_doc=export_doc,
                    export_images=export_images,
                    export_json=export_json,
                    provider_name=(session.selected_provider or session.provider_match.name or "").strip(),
                )
                files_processed += 1
                total_images += int(result.get("image_count", 0) or 0)
            except Exception as exc:
                failures.append(f"{Path(session.path).name}: {exc}")

        if export_doc:
            self.set_status(f"Completed DOC export for {files_processed} files.")
        elif export_images and not export_json:
            self.set_status(f"Completed Image export of {total_images} images from {files_processed} files.")
        elif export_json and not export_doc:
            self.set_status(f"Completed JSON export for {files_processed} files.")
        else:
            self.set_status("Batch export complete")

        if failures:
            messagebox.showerror(APP_TITLE, "Batch export encountered errors:\n\n" + "\n".join(failures[:20]))

    def export_json_string(self):
        if not self.current_file:
            messagebox.showerror(APP_TITLE, "Please drag in a PDF, DOCX, DOC, EML, or MSG first.")
            return
        values = self.prepare_export_values(self.collect_values())
        # Silently no-op if Work Provider is empty — by design we never
        # produce an "UnknownVRM.json" for an unidentified document.
        if not (values.get("work_provider") or "").strip():
            return
        result = self.export_outputs_for_values(
            source_file=self.current_source_files or [self.current_file],
            values=values,
            export_doc=False,
            export_images=False,
            export_json=True,
            provider_name=(self.selected_provider_var.get() or self.detected_provider_var.get() or "").strip(),
        )
        json_string = str(result.get("json_string", ""))

        try:
            self.root.clipboard_clear()
            self.root.clipboard_append(json_string)
        except Exception:
            pass

        self.set_status("Exported JSON to Desktop")

    def load_provider_into_editor(self, provider_config: dict):
        provider_config = self.engine.normalize_provider_config(provider_config)
        self.selected_provider_var.set(provider_config.get("name", ""))
        self.provider_name_var.set(provider_config.get("name", ""))

        self.detect_box.delete("1.0", tk.END)
        phrases = provider_config.get("detect_phrases", [])
        if phrases:
            self.detect_box.insert("1.0", "\n".join(phrases))

        field_rules = provider_config.get("field_rules", {})
        self.use_current_date_var.set(bool(provider_config.get("use_current_date_for_inspection_date")))
        self.force_postcode_var.set(bool(provider_config.get("force_postcode_for_inspection_address")))
        self.engineer_report_var.set(bool(provider_config.get("engineer_report")))

        # Work Provider field — manual-input only, no method dropdown.
        wp_rule = field_rules.get("work_provider") or {}
        wp_value = str((wp_rule or {}).get("config", "") or "")
        self.mapping_config_vars["work_provider"].set(wp_value)

        for key in NON_PROVIDER_FIELDS:
            rule = field_rules.get(key, {})
            method = rule.get("method") or DEFAULT_METHOD_BY_FIELD.get(key, "single_label")
            config = rule.get("config", "")
            display_method = LEGACY_METHOD_TO_DISPLAY_CODE.get(method, method)

            # Presence-check fields (VAT Status, Mileage Unit) are always
            # displayed as single_label. Legacy presets may have used
            # two_labels with "foo || bar" syntax; collapse that to just
            # the start label, which was the positive-scenario token under
            # the old logic.
            if key in PRESENCE_CHECK_FIELDS:
                if display_method == "two_labels" and "||" in str(config):
                    start_label, _ = self.engine.parse_two_label_config(config)
                    config = start_label
                self.mapping_method_vars[key].set(METHOD_CODE_TO_LABEL["single_label"])
                self.mapping_config_vars[key].set(config or "")
                self.mapping_second_config_vars[key].set("")
                self.update_method_ui(key)
                continue

            self.mapping_method_vars[key].set(METHOD_CODE_TO_LABEL.get(display_method, METHOD_CODE_TO_LABEL["single_label"]))
            if key == "inspection_date" and str(config).strip().lower() == "{today}" and not provider_config.get("use_current_date_for_inspection_date"):
                self.use_current_date_var.set(True)
                self.mapping_config_vars[key].set("")
                self.mapping_second_config_vars[key].set("")
            elif display_method in ("two_labels", "fixed_position_label", "single_label_offset"):
                start_label, end_label = self.engine.parse_two_label_config(config)
                self.mapping_config_vars[key].set(start_label)
                self.mapping_second_config_vars[key].set(end_label)
            else:
                self.mapping_config_vars[key].set(config)
                self.mapping_second_config_vars[key].set("")
            self.update_method_ui(key)

    def clear_provider_editor(self, keep_provider_dropdown: bool = False):
        if not keep_provider_dropdown:
            self.selected_provider_var.set("")
        self.provider_name_var.set("")
        self.detect_box.delete("1.0", tk.END)
        self.use_current_date_var.set(False)
        self.force_postcode_var.set(False)
        self.engineer_report_var.set(False)
        self.mapping_config_vars["work_provider"].set("")
        for key in NON_PROVIDER_FIELDS:
            self.mapping_method_vars[key].set(METHOD_CODE_TO_LABEL[DEFAULT_METHOD_BY_FIELD.get(key, "single_label")])
            self.mapping_config_vars[key].set("")
            self.mapping_second_config_vars[key].set("")
            self.update_method_ui(key)

    def on_provider_selected(self, _event=None):
        name = self.selected_provider_var.get().strip()
        if not name:
            return
        provider = self.engine.get_provider_config(name)
        if provider:
            self.load_provider_into_editor(provider)
            self.set_status(f"Ready to use provider: {name}")

    def use_selected_provider(self):
        if not self.current_text:
            messagebox.showinfo(APP_TITLE, "Please drag in a PDF, DOCX, DOC, EML, or MSG first.")
            return

        selected_name = self.selected_provider_var.get().strip()
        provider_config = self.engine.get_provider_config(selected_name) if selected_name else None

        if provider_config:
            provider = ProviderMatch(selected_name, 999, provider_config)
            self.rescan_with_provider(provider)
            self.save_current_batch_state()
            return

        if self.current_provider_match and self.current_provider_match.name != "Unknown / Unmapped":
            self.rescan_with_provider(self.current_provider_match)
            self.save_current_batch_state()
            return

        guessed = self.get_field_value("work_provider").strip()
        self.clear_provider_editor(keep_provider_dropdown=True)
        if guessed:
            self.provider_name_var.set(guessed)
        self.set_status("No saved provider selected. Enter a new provider mapping and save it.")

    def save_provider_mapping(self):
        # Provider Name is the dropdown label (e.g. "FW (Garage)").
        # If the user didn't type a Provider Name, fall back to the
        # Work Provider value they typed in the rules table — for the
        # common case of a single-format provider where the dropdown
        # name and the JSON output value are the same string.
        provider_name = self.provider_name_var.get().strip() or self.mapping_config_vars["work_provider"].get().strip()
        if not provider_name:
            messagebox.showinfo(APP_TITLE, "Please enter a provider name before saving.")
            return

        detect_phrases = [line.strip() for line in self.detect_box.get("1.0", tk.END).splitlines() if line.strip()]
        field_rules = {}

        # Work Provider is a manual-input-only field. Its rule is
        # ``{"method": "manual_input", "config": "<user text>"}``.
        field_rules["work_provider"] = {
            "method": "manual_input",
            "config": self.mapping_config_vars["work_provider"].get().strip(),
        }

        for key in NON_PROVIDER_FIELDS:
            method_label = self.mapping_method_vars[key].get().strip() or METHOD_CODE_TO_LABEL[DEFAULT_METHOD_BY_FIELD.get(key, "single_label")]
            method = METHOD_LABEL_TO_CODE.get(method_label, DEFAULT_METHOD_BY_FIELD.get(key, "single_label"))

            # Presence-check fields don't use a method — always save as
            # single_label with the raw token list as config.
            if key in PRESENCE_CHECK_FIELDS:
                method = "single_label"
                config_value = self.mapping_config_vars[key].get().strip()
            elif method in ("two_labels", "fixed_position_label", "single_label_offset"):
                start_label = self.mapping_config_vars[key].get().strip()
                end_label = self.mapping_second_config_vars[key].get().strip()
                config_value = f"{start_label} || {end_label}" if (start_label or end_label) else ""
            else:
                config_value = self.mapping_config_vars[key].get().strip()
            field_rules[key] = {
                "method": method,
                "config": config_value,
            }

        provider_data = {
            "name": provider_name,
            "detect_phrases": detect_phrases,
            "field_rules": field_rules,
            "use_current_date_for_inspection_date": bool(self.use_current_date_var.get()),
            "force_postcode_for_inspection_address": bool(self.force_postcode_var.get()),
            "engineer_report": bool(self.engineer_report_var.get()),
        }

        self.engine.upsert_provider(provider_data)
        self.refresh_provider_dropdown()
        self.selected_provider_var.set(provider_name)
        self.set_field_value("work_provider", provider_name)
        self.set_status(f"Saved mapping for {provider_name}")

    def delete_provider_mapping(self):
        selected_name = self.selected_provider_var.get().strip() or self.provider_name_var.get().strip()
        if not selected_name:
            messagebox.showinfo(APP_TITLE, "Please choose a saved provider to delete.")
            return

        provider = self.engine.get_provider_config(selected_name)
        if not provider:
            messagebox.showinfo(APP_TITLE, f'No saved mapping found for "{selected_name}".')
            return

        if not messagebox.askyesno(APP_TITLE, f'Delete the mapping for "{selected_name}"?'):
            return

        providers = [
            p for p in self.engine.config.get("providers", [])
            if p.get("name", "").strip().lower() != selected_name.lower()
        ]
        self.engine.config["providers"] = providers
        self.engine.save_config()
        self.refresh_provider_dropdown()
        self.clear_provider_editor()
        if self.current_provider_match and self.current_provider_match.name.lower() == selected_name.lower():
            self.detected_provider_var.set("Unknown / Unmapped")
        self.set_status(f"Deleted mapping for {selected_name}")


    def _set_value_widget_override_border(self, key: str, active: bool):
        """Show or hide the red 'engineer-overwritten' border on a single
        value entry widget on the left panel.

        Implemented via Tk's highlight ring (``highlightthickness`` +
        ``highlightbackground``) so it surrounds the widget without altering
        its existing relief or contents.
        """
        widget = self.value_entry_widgets.get(key)
        if widget is None:
            return
        try:
            if active:
                widget.configure(
                    highlightthickness=2,
                    highlightbackground=ENGINEER_OVERRIDE_BORDER_COLOR,
                    highlightcolor=ENGINEER_OVERRIDE_BORDER_COLOR,
                )
            else:
                widget.configure(
                    highlightthickness=0,
                    highlightbackground="white",
                    highlightcolor="white",
                )
        except Exception:
            pass

    def apply_engineer_override_highlights(self):
        """Refresh the red border on every value entry to reflect the
        current ``engineer_overridden_fields`` set."""
        for key in FIELD_KEYS:
            self._set_value_widget_override_border(key, key in self.engineer_overridden_fields)

    def clear_engineer_override_highlights(self):
        """Drop all engineer-override state and remove the red borders.
        Called whenever a fresh document is dragged in."""
        self.engineer_overridden_fields = set()
        self.apply_engineer_override_highlights()

    @staticmethod
    def _compute_engineer_overrides(base_values: Dict[str, str], engineer_values: Dict[str, str]) -> set:
        """Return the set of field keys that an engineer report has touched.

        A field counts as 'overridden' whenever the engineer report
        produces a non-blank value for it — regardless of whether the
        value differs from what was already there. Highlighting the
        field even when the value matches confirms to the user that
        their engineer-report mapping rule is firing as intended.

        Mirrors the merge logic used in ``apply_engineer_report_session``
        and ``combine_instruction_and_engineer_report``: those functions
        write any non-blank engineer value into the merged result, and
        this set is the visual record of which fields they touched.
        """
        # ``base_values`` is no longer consulted but kept in the signature
        # so callers don't have to change.
        del base_values
        overrides: set = set()
        for key in NON_PROVIDER_FIELDS:
            new_value = (engineer_values.get(key) or "").strip()
            if new_value:
                overrides.add(key)
        return overrides

    def clear_field_highlights(self):
        if not self.preview_widget:
            return
        for field_name in self.field_highlight_colors:
            self.preview_widget.tag_remove(f"field_{field_name}", "1.0", tk.END)

    def highlight_vrm_in_preview(self, value: str):
        if not self.preview_widget:
            return
        normalized_target = normalize_vrm_value(value)
        if len(normalized_target) < 3:
            return

        total_lines = int(float(self.preview_widget.index("end-1c")))
        prefix_pattern = re.compile(r"^\s*\d+\s*\|\s?")

        for line_no in range(1, total_lines + 1):
            line_text = self.preview_widget.get(f"{line_no}.0", f"{line_no}.end")
            if not line_text:
                continue

            prefix_match = prefix_pattern.match(line_text)
            prefix_len = len(prefix_match.group(0)) if prefix_match else 0
            content = line_text[prefix_len:]

            normalized_chars = []
            char_positions = []
            for idx, ch in enumerate(content):
                if not ch.isspace():
                    normalized_chars.append(ch.upper())
                    char_positions.append(idx)

            normalized_content = "".join(normalized_chars)
            if not normalized_content:
                continue

            search_from = 0
            while True:
                match_index = normalized_content.find(normalized_target, search_from)
                if match_index == -1:
                    break
                start_char = char_positions[match_index]
                end_char = char_positions[match_index + len(normalized_target) - 1] + 1
                self.preview_widget.tag_add(
                    "field_vrm",
                    f"{line_no}.{prefix_len + start_char}",
                    f"{line_no}.{prefix_len + end_char}",
                )
                search_from = match_index + 1

    def highlight_extracted_values(self, values: Dict[str, str]):
        if not self.preview_widget:
            return
        self.clear_field_highlights()
        for field_name, raw_value in values.items():
            value = (raw_value or "").strip()
            if not value:
                continue
            if field_name == "vrm":
                self.highlight_vrm_in_preview(value)
                continue
            snippets = [part.strip() for part in value.splitlines() if part.strip()]
            if not snippets:
                snippets = [value]
            for snippet in snippets:
                if len(snippet) < 3:
                    continue
                start = "1.0"
                while True:
                    idx = self.preview_widget.search(snippet, start, stopindex=tk.END, nocase=True)
                    if not idx:
                        break
                    end = f"{idx}+{len(snippet)}c"
                    self.preview_widget.tag_add(f"field_{field_name}", idx, end)
                    start = end

    def copy_preview_selection(self, _event=None):
        if not self.preview_widget:
            return "break"
        try:
            selected = self.preview_widget.get(tk.SEL_FIRST, tk.SEL_LAST)
        except tk.TclError:
            return None
        cleaned_lines = [re.sub(r"^\s*\d+\s*\|\s?", "", line) for line in selected.splitlines()]
        cleaned = "\n".join(cleaned_lines)
        self.root.clipboard_clear()
        self.root.clipboard_append(cleaned)
        return "break"

    def focus_find(self, _event=None):
        self.find_entry.focus_set()
        self.find_entry.selection_range(0, tk.END)
        return "break"

    def clear_find_tags(self):
        if self.preview_widget:
            self.preview_widget.tag_remove("find_all", "1.0", tk.END)
            self.preview_widget.tag_remove("find_current", "1.0", tk.END)
        self.find_matches = []
        self.find_index = -1
        self.last_find_term = ""

    def update_find_matches(self):
        if not self.preview_widget:
            return
        term = self.find_var.get()
        self.clear_find_tags()
        if not term:
            return
        self.last_find_term = term
        start = "1.0"
        while True:
            idx = self.preview_widget.search(term, start, stopindex=tk.END, nocase=True)
            if not idx:
                break
            end = f"{idx}+{len(term)}c"
            self.preview_widget.tag_add("find_all", idx, end)
            self.find_matches.append(idx)
            start = end

    def select_find_match(self, step: int):
        if not self.preview_widget:
            return "break"
        term = self.find_var.get().strip()
        if not term:
            return "break"
        if term != self.last_find_term or not self.find_matches:
            self.update_find_matches()
        if not self.find_matches:
            self.set_status(f'No matches for "{term}"')
            return "break"

        self.find_index = (self.find_index + step) % len(self.find_matches)
        idx = self.find_matches[self.find_index]
        end = f"{idx}+{len(term)}c"
        self.preview_widget.tag_remove("find_current", "1.0", tk.END)
        self.preview_widget.tag_add("find_current", idx, end)
        self.preview_widget.mark_set(tk.INSERT, idx)
        self.preview_widget.see(idx)
        self.set_status(f'Match {self.find_index + 1} of {len(self.find_matches)} for "{term}"')
        return "break"

    def find_next(self, _event=None):
        return self.select_find_match(1)

    def find_previous(self, _event=None):
        return self.select_find_match(-1)


def main():
    # Try to point pytesseract at a bundled Tesseract binary if one is
    # present alongside the .exe. Silently no-ops if Tesseract isn't
    # bundled — the OCR fallback will simply not fire.
    configure_bundled_tesseract()
    root = TkinterDnD.Tk() if HAS_DND else tk.Tk()
    apply_window_icon(root)
    style = ttk.Style(root)
    try:
        style.theme_use("clam")
    except Exception:
        pass
    App(root)
    root.mainloop()


if __name__ == "__main__":
    main()